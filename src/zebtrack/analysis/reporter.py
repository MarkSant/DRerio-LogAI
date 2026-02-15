"""Report generation module for behavioral analysis results.

Generates Word document reports with visualizations and metrics from zebrafish
tracking data, supporting both individual and project-level reports.
"""

import gettext
import io
import locale
import os
import tempfile
import warnings
from collections.abc import Callable, Sequence
from pathlib import Path
from typing import Any

import pandas as pd
import structlog
from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Inches
from docxtpl import DocxTemplate

from zebtrack.analysis.analysis_service import AnalysisService
from zebtrack.analysis.data_transformer import DISPLAY_COLUMN_MAPPING, DataTransformer
from zebtrack.analysis.models import AnalysisResult
from zebtrack.analysis.roi import ROI
from zebtrack.analysis.visualization_generator import VisualizationGenerator

log = structlog.get_logger(__name__)

_BASE_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = _BASE_DIR.parent / "templates"
LOCALES_DIR = _BASE_DIR.parent / "locales"
REPORTER_DOMAIN = "reporter"
INDIVIDUAL_REPORT_TEMPLATE = TEMPLATES_DIR / "individual_report_template.docx"
PROJECT_REPORT_TEMPLATE = TEMPLATES_DIR / "project_report_template.docx"


def _load_translator():
    languages: list[str] = []

    env_candidates: list[str] = []
    for env_var in ("LANGUAGE", "LC_ALL", "LC_MESSAGES", "LANG"):
        value = os.environ.get(env_var)
        if value:
            env_candidates.extend(lang for lang in value.split(":") if lang)

    try:
        current_locale = locale.getlocale()[0]
    except (AttributeError, TypeError, ValueError):
        current_locale = None

    for candidate in [*env_candidates, current_locale]:
        if not candidate:
            continue
        if candidate not in languages:
            languages.append(candidate)
        if "_" in candidate:
            base = candidate.split("_", 1)[0]
            if base not in languages:
                languages.append(base)

    try:
        translator = gettext.translation(
            REPORTER_DOMAIN,
            localedir=str(LOCALES_DIR),
            languages=languages if languages else None,
            fallback=True,
        )
    except OSError:
        translator = gettext.NullTranslations()

    return translator.gettext


_translator: Callable[[str], str] = _load_translator()


def _(message: str) -> str:
    return _translator(message)


def _format_time_minutes_seconds(seconds):
    """Format time as MM:SS or HH:MM:SS if over an hour."""
    if pd.isna(seconds) or seconds is None:
        return "N/A"

    seconds = float(seconds)
    if seconds < 60:
        return f"{int(seconds)}s"
    elif seconds < 3600:
        minutes = int(seconds // 60)
        secs = int(seconds % 60)
        return f"{minutes}:{secs:02d}"
    else:
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        return f"{hours}:{minutes:02d}:{secs:02d}"


class Reporter:
    """Generate analysis reports in various formats (Excel, Word).

    This class is responsible for the final assembly of reports, delegating
    data transformation to DataTransformer and visualization generation to
    VisualizationGenerator.

    Refactored in Phase 2 (Task 2.5) to separate concerns:
        - DataTransformer: Handles data transformation and tidying
        - VisualizationGenerator: Handles plot/visualization generation
        - Reporter: Handles final report assembly and export

    Example:
        >>> # Modern approach (RECOMMENDED)
        >>> service = AnalysisService(settings_obj=settings)
        >>> analysis = service.run_full_analysis_as_dto(...)
        >>> reporter = Reporter.from_analysis(analysis)
        >>> reporter.export_summary_data("output.xlsx")
        >>> reporter.export_individual_report("report.docx")
    """

    def __init__(
        self,
        trajectory_df: pd.DataFrame = None,
        metadata: dict | None = None,
        # Calibration and setup
        pixelcm_x: float | None = None,
        pixelcm_y: float | None = None,
        video_height_px: int | None = None,
        arena_polygon_px: Sequence[Sequence[float]] | None = None,
        rois: list[ROI] | None = None,
        fps: float | None = None,
        # Optional params
        roi_colors: dict | None = None,
        video_path: str | None = None,
        calibration=None,
        frame_crop_box: tuple[int, int, int, int] | None = None,
        # Analysis params
        sharp_turn_threshold: float = 90.0,
        freezing_threshold: float = 1.5,
        freezing_duration: float = 1.0,
        smoothing_window_length: int | None = None,
        smoothing_polyorder: int | None = None,
        settings_obj=None,
        behavioral_config: dict | None = None,
        # Modern path: DTO-based construction
        analysis: AnalysisResult | None = None,
    ):
        """Create Reporter instance for generating analysis reports.

        **RECOMMENDED**: Use Reporter.from_analysis(analysis_result) instead
        of calling this constructor directly.

        Construction Paths:
            1. Modern (RECOMMENDED): Pass pre-computed analysis via `analysis` parameter
            2. Legacy (DEPRECATED): Pass trajectory_df + all other parameters

        Args:
            trajectory_df: Raw trajectory DataFrame (DEPRECATED - use analysis parameter)
            metadata: Experiment metadata (DEPRECATED - included in analysis)
            pixelcm_x: Pixels-to-cm conversion (DEPRECATED - included in analysis)
            pixelcm_y: Pixels-to-cm conversion (DEPRECATED - included in analysis)
            video_height_px: Video height (DEPRECATED - included in analysis)
            arena_polygon_px: Arena polygon (DEPRECATED - included in analysis)
            rois: ROI list (DEPRECATED - included in analysis)
            fps: Frame rate (DEPRECATED - included in analysis)
            roi_colors: ROI colors (optional)
            video_path: Video file path (optional)
            calibration: Calibration object (optional)
            sharp_turn_threshold: Sharp turn threshold (deg/s)
            freezing_threshold: Freezing velocity threshold (cm/s)
            freezing_duration: Minimum freezing duration (s)
            smoothing_window_length: Smoothing window (optional)
            smoothing_polyorder: Smoothing polynomial order (optional)
            settings_obj: Settings instance (optional)
            analysis: AnalysisResult DTO (RECOMMENDED - modern path)

        Example (Modern - RECOMMENDED):
            >>> # Get pre-computed analysis
            >>> service = AnalysisService(settings_obj=settings)
            >>> analysis = service.run_full_analysis_as_dto(...)
            >>> reporter = Reporter.from_analysis(analysis)  # Use factory method!

        Example (Legacy - DEPRECATED):
            >>> reporter = Reporter(
            ...     trajectory_df=df,
            ...     metadata={"experiment_id": "exp_001"},
            ...     pixelcm_x=10.0,
            ...     # ... many more parameters ...
            ... )
            # DeprecationWarning: Use Reporter.from_analysis() instead
        """
        # Modern path: using AnalysisResult DTO
        if analysis is not None:
            # Delegate to from_analysis() and copy attributes
            temp_instance = Reporter.from_analysis(analysis)
            self.__dict__.update(temp_instance.__dict__)
            return

        # Legacy path validation
        if trajectory_df is None:
            raise ValueError(
                "Reporter: Either 'analysis' or 'trajectory_df' must be provided. "
                "RECOMMENDED: Use Reporter.from_analysis(analysis_result) instead."
            )

        # Emit deprecation warning for legacy constructor
        warnings.warn(
            "Reporter: Direct instantiation with trajectory_df is DEPRECATED and "
            "will be removed in v3.0. "
            "\n"
            "Migration Guide:\n"
            "  Instead of: Reporter(trajectory_df=df, metadata=meta, pixelcm_x=10.0, ...)\n"
            "  Use:        service = AnalysisService(settings_obj=settings)\n"
            "              analysis = service.run_full_analysis_as_dto(...)\n"
            "              reporter = Reporter.from_analysis(analysis)\n"
            "\n"
            "Benefits: Better performance (no re-analysis), improved testability, type safety.\n"
            "Timeline: Deprecation in v2.1 (2025-11), removal in v3.0 (2026-02).\n"
            "See docs/migration/reporter-v3.md for details.",
            DeprecationWarning,
            stacklevel=2,
        )

        # Legacy constructor logic
        self.settings = settings_obj
        self.metadata = metadata
        self.roi_colors = roi_colors if roi_colors else {}
        self.video_path = video_path
        self.calibration = calibration
        self.frame_crop_box = frame_crop_box
        self._pixelcm_x = pixelcm_x
        self._pixelcm_y = pixelcm_y
        self._video_height_px = video_height_px

        # Ensure trajectory coordinates stay aligned with calibration space
        # before running any calculations.
        if calibration is not None:
            trajectory_df = DataTransformer.warp_trajectory_if_needed(
                trajectory_df, calibration, force=True
            )

        if "track_id" in trajectory_df.columns:
            track_ids = pd.to_numeric(trajectory_df["track_id"], errors="coerce")
            trajectory_df = trajectory_df.copy()
            trajectory_df["track_id"] = track_ids.astype("Int64")

        # Run the unified analysis via the service
        service = AnalysisService(settings_obj=settings_obj)
        (
            self.report,
            self.b_analyzer,
            self.r_analyzer,
            self.validation_warnings,
            self.validation_stats,
        ) = service.run_full_analysis(
            trajectory_df=trajectory_df,
            pixelcm_x=pixelcm_x or 1.0,
            pixelcm_y=pixelcm_y or 1.0,
            video_height_px=video_height_px or 0,
            arena_polygon_px=arena_polygon_px or [],
            rois=rois or [],
            fps=fps or 30.0,
            freezing_vel_threshold=freezing_threshold,
            freezing_min_duration=freezing_duration,
            smoothing_window_length=smoothing_window_length,
            smoothing_polyorder=smoothing_polyorder,
            behavioral_config=behavioral_config,
        )

        # Store for plotting methods that still need them
        self.sharp_turn_threshold = sharp_turn_threshold
        self.freezing_threshold = freezing_threshold
        self.freezing_duration = freezing_duration
        self.validation_warnings = self.report.get("validacao", {}).get("avisos", [])
        self.validation_stats = self.report.get("validacao", {}).get("estatisticas", {})

        # Initialize data transformer and visualization generator
        self.data_transformer = DataTransformer()
        # Store behavioral_config from parameter (critical for geotaxis zone calculations)
        # This MUST be set before create_tidy_dataframe is called
        self.behavioral_config = behavioral_config if behavioral_config else {}

        self.viz_generator = VisualizationGenerator(
            b_analyzer=self.b_analyzer,
            r_analyzer=self.r_analyzer,
            metadata=self.metadata or {},
            roi_colors=self.roi_colors,
            calibration=self.calibration,
            pixelcm_x=self._pixelcm_x,
            pixelcm_y=self._pixelcm_y,
            video_height_px=self._video_height_px,
            sharp_turn_threshold=self.sharp_turn_threshold,
            settings_obj=self.settings,
            frame_crop_box=self.frame_crop_box,
            behavioral_config=self.behavioral_config,
        )

        # Generate the tidy dataframe from the report
        tidy_df = self.data_transformer.create_tidy_dataframe(
            report=self.report,
            metadata=self.metadata or {},
            b_analyzer=self.b_analyzer,
            r_analyzer=self.r_analyzer,
            roi_colors=self.roi_colors,
            validation_stats=self.validation_stats,
            behavioral_config=getattr(self, "behavioral_config", {}),
        )
        self.tidy_data = self.data_transformer.standardize_tidy_dataframe(
            tidy_df, self.metadata or {}
        )

    @classmethod
    def from_analysis(cls, analysis: AnalysisResult) -> "Reporter":
        """Create Reporter from pre-computed instance (RECOMMENDED)."""
        instance = cls.__new__(cls)

        # Store settings reference
        instance.settings = None

        # Store metadata and calibration
        instance.metadata = analysis.metadata
        instance.roi_colors = analysis.roi_colors
        instance.video_path = analysis.video_path
        instance.calibration = analysis.calibration_params.calibration
        instance.frame_crop_box = getattr(analysis, "frame_crop_box", None)
        instance._pixelcm_x = analysis.calibration_params.pixelcm_x
        instance._pixelcm_y = analysis.calibration_params.pixelcm_y
        instance._video_height_px = analysis.calibration_params.video_height_px

        # Store analysis results directly
        instance.report = analysis.report
        instance.b_analyzer = analysis.behavioral_analyzer
        instance.r_analyzer = analysis.roi_analyzer
        instance.behavioral_config = analysis.behavioral_config or {}

        # Store analysis parameters
        instance.sharp_turn_threshold = analysis.sharp_turn_threshold
        instance.freezing_threshold = analysis.freezing_threshold
        instance.freezing_duration = analysis.freezing_duration
        instance.validation_warnings = getattr(analysis, "validation_warnings", [])
        instance.validation_stats = getattr(analysis, "validation_stats", {})

        # Invoke common initialization logic (creates viz generator and tidy df)
        # Note: We must verify if __init__ logic can be reused or must be duplicated.
        # Looking at original code, Step 294 showed __init__ ending with viz_generator creation
        # but from_analysis was separate.
        # Actually, from_analysis calls instance.__init__? No, current impl calls __new__.
        # So we need to call the setup logic manually.

        # Initialize components (Duplicated or extracted? Let's check context)
        # Re-using the logic block from __init__ (lines 278-302) via a helper would be better
        # but for now we duplicate or invoke __init__ logic if it was in a method.
        # It seems lines 278-302 logic IS part of __init__.
        # We should probably manually trigger lines 278-302 logic here.

        instance.data_transformer = DataTransformer()
        instance.viz_generator = VisualizationGenerator(
            b_analyzer=instance.b_analyzer,
            r_analyzer=instance.r_analyzer,
            metadata=instance.metadata,
            roi_colors=instance.roi_colors,
            calibration=instance.calibration,
            pixelcm_x=instance._pixelcm_x,
            pixelcm_y=instance._pixelcm_y,
            video_height_px=instance._video_height_px,
            sharp_turn_threshold=instance.sharp_turn_threshold,
            settings_obj=instance.settings,
            frame_crop_box=instance.frame_crop_box,
            behavioral_config=instance.behavioral_config,
        )

        tidy_df = instance.data_transformer.create_tidy_dataframe(
            report=instance.report,
            metadata=instance.metadata,
            b_analyzer=instance.b_analyzer,
            r_analyzer=instance.r_analyzer,
            roi_colors=instance.roi_colors,
            validation_stats=instance.validation_stats,
            behavioral_config=instance.behavioral_config,
        )
        instance.tidy_data = instance.data_transformer.standardize_tidy_dataframe(
            tidy_df, instance.metadata
        )

        return instance

        # Initialize data transformer and visualization generator
        instance.data_transformer = DataTransformer()
        instance.viz_generator = VisualizationGenerator(
            b_analyzer=instance.b_analyzer,
            r_analyzer=instance.r_analyzer,
            metadata=instance.metadata,
            roi_colors=instance.roi_colors,
            calibration=instance.calibration,
            pixelcm_x=instance._pixelcm_x,
            pixelcm_y=instance._pixelcm_y,
            video_height_px=instance._video_height_px,
            sharp_turn_threshold=instance.sharp_turn_threshold,
            settings_obj=instance.settings,
            frame_crop_box=instance.frame_crop_box,
        )

        # Generate tidy dataframe
        tidy_df = instance.data_transformer.create_tidy_dataframe(
            report=instance.report,
            metadata=instance.metadata,
            b_analyzer=instance.b_analyzer,
            r_analyzer=instance.r_analyzer,
            roi_colors=instance.roi_colors,
            validation_stats=instance.validation_stats,
            behavioral_config=analysis.behavioral_config,
        )
        instance.tidy_data = instance.data_transformer.standardize_tidy_dataframe(
            tidy_df, instance.metadata
        )

        return instance

    def export_summary_data(
        self,
        output_path: Path | str,
        format: str = "excel",
        expected_roi_names: list[str] | None = None,
    ):
        """Export summary data to file (Excel, CSV, or Parquet).

        Args:
            output_path: Output file path
            format: Output format ('excel', 'csv', or 'parquet')
            expected_roi_names: Optional list of ROI names for schema standardization.
                If provided, missing ROI columns will be added with NaN/0 defaults.

        Raises:
            ValueError: If unsupported format is specified
        """
        path = Path(output_path) if isinstance(output_path, str) else output_path
        path.parent.mkdir(parents=True, exist_ok=True)

        # Standardize ROI columns if expected names provided
        data_to_export = self.tidy_data
        if expected_roi_names:
            data_to_export = self.data_transformer.standardize_roi_columns(
                self.tidy_data, expected_roi_names
            )

        self.data_transformer.validate_schema(data_to_export)

        # Apply Geotaxis column renaming if applicable
        # Requires behavioral info which might be in metadata or config
        height_cm = float((self.metadata or {}).get("aquarium_height_cm", 0) or 0)
        num_zones = int((self.metadata or {}).get("geotaxis_num_zones", 0) or 0)

        # Try to fetch from behavioral config if available
        if num_zones == 0 and hasattr(self, "behavioral_config"):
            num_zones = int(self.behavioral_config.get("geotaxis_num_zones", 0) or 0)

        if format in ("excel", "csv"):
            # Apply display formatting (renaming + geotaxis) for human-readable formats
            data_to_export = self.data_transformer.prepare_for_display(
                data_to_export, height_cm, num_zones
            )
        # For parquet, we keep internal names (do nothing)

        if format == "excel":
            data_to_export.to_excel(output_path, index=False, engine="openpyxl")
        elif format == "csv":
            data_to_export.to_csv(output_path, index=False)
        elif format == "parquet":
            data_to_export.to_parquet(output_path, index=False)
        else:
            raise ValueError(f"Unsupported file format: {format}")

    def export_individual_report(self, output_path: Path | str):
        """Export a complete individual report.

        This is a convenience wrapper around the step-by-step method for consumers
        who don't need progress updates.

        Args:
            output_path: Output file path for the report.
        """
        self.export_individual_report_step_by_step(output_path, lambda p, s: None)

    def export_individual_report_step_by_step(
        self, output_path: Path | str, progress_callback: Callable[[float, str], None]
    ):
        """Export individual report with step-by-step progress reporting.

        Args:
            output_path: Output file path for the report
            progress_callback: Callback function (progress: float, status: str) -> None
        """
        output_path = Path(output_path) if isinstance(output_path, str) else output_path
        total_steps = 11
        template_path = INDIVIDUAL_REPORT_TEMPLATE
        heading_text = _("Analysis Report - {experiment_id}").format(
            experiment_id=(self.metadata or {}).get("experiment_id", "Unknown")
        )

        doc_template, document = self._prepare_report_document(template_path, heading_text)

        # Step 1: Metadata section
        self._append_metadata_section(document, progress_callback, total_steps)
        # Final step: ensure document is written to disk. If a DocxTemplate
        # was used, save via its API to preserve any template context. Otherwise
        # save the Document object directly.
        try:
            file_path = (
                f"{output_path}.docx"
                if not str(output_path).lower().endswith(".docx")
                else str(output_path)
            )
            if doc_template is not None:
                # DocxTemplate saves via .save()
                doc_template.save(file_path)
            else:
                document.save(file_path)
            progress_callback(1.0, _("Report saved"))
            log.info("reporter.individual_report.saved", path=file_path)
        except Exception as e:
            log.error("reporter.individual_report.save_failed", error=str(e), exc_info=True)

    def _prepare_report_document(self, template_path: Path, heading_text: str):
        """Prepare a docx document using template if available, return (doc_template, document)."""
        doc_template: DocxTemplate | None = None
        document: DocxDocument | None = None
        template_rendered = False

        if template_path.exists():
            try:
                doc_template = DocxTemplate(str(template_path))
                doc_template.render({"title": heading_text})
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
                    tmp_path = Path(tmp_file.name)

                try:
                    doc_template.save(str(tmp_path))
                    document = Document(str(tmp_path))
                    template_rendered = True
                finally:
                    tmp_path.unlink(missing_ok=True)

                # Continue using python-docx document API for append operations
                doc_template = None
            except Exception:
                log.warning(
                    "analysis.reporter.template_render_failed",
                    template=str(template_path),
                    exc_info=True,
                )
                document = Document()
        else:
            log.warning(
                "analysis.reporter.template_missing_fallback",
                template=str(template_path),
            )
            document = Document()

        if document is None:
            document = Document()

        if doc_template is None and not template_rendered:
            document.add_heading(heading_text, level=1)

        return doc_template, document

    def _append_metadata_section(
        self,
        document: DocxDocument,
        progress_callback: Callable[[float, str], None],
        total_steps: int,
    ) -> None:
        """Append the metadata section to the provided document and call progress callback."""
        document.add_heading(_("Experiment Metadata"), level=2)
        for key, value in (self.metadata or {}).items():
            document.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
        progress_callback(1 / total_steps, _("Metadata added"))

        # Step 2: Summary table
        document.add_heading(_("Metrics Summary"), level=2)
        df = self.tidy_data.drop(
            columns=[k for k in (self.metadata or {}).keys() if k in self.tidy_data.columns]
        )

        # Apply Geotaxis column renaming for Word report
        height_cm = float((self.metadata or {}).get("aquarium_height_cm", 0) or 0)
        num_zones = int((self.metadata or {}).get("geotaxis_num_zones", 0) or 0)
        if hasattr(self, "behavioral_config") and num_zones == 0:
            num_zones = int(self.behavioral_config.get("geotaxis_num_zones", 0) or 0)

        if height_cm > 0 and num_zones > 0:
            df = self.data_transformer.rename_geotaxis_columns(df, height_cm, num_zones)
        else:
            # Fallback: rename geotaxis columns generically if dimensions unavailable
            # Use 1-indexed zone numbers for user display (Zone 0 internal -> Zone 1 display)
            rename_geo = {}
            for col in df.columns:
                if col.startswith("geotaxis_zone_") and col.endswith("_pct"):
                    try:
                        idx = int(col.split("_")[2])
                        # Zone 0 = bottom, display as "Zona 1 (Fundo)"
                        if idx == 0:
                            rename_geo[col] = "Geotaxis Zona 1 - Fundo (%)"
                        else:
                            rename_geo[col] = f"Geotaxis Zona {idx + 1} (%)"
                    except (IndexError, ValueError):
                        log.debug("reporter.geotaxis_rename.parse_error", column=col, exc_info=True)
            if rename_geo:
                df = df.rename(columns=rename_geo)

        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.cell(0, 0).text = _("Metric")
        table.cell(0, 1).text = _("Value")

        for column_name in df.columns:
            row_cells = table.add_row().cells
            # Use DISPLAY_COLUMN_MAPPING for proper formatting (e.g., "Max Speed (cm/s)")
            # Fall back to title case only if no mapping exists
            formatted_name = DISPLAY_COLUMN_MAPPING.get(
                column_name, column_name.replace("_", " ").title()
            )
            row_cells[0].text = formatted_name

            value = df[column_name].iloc[0]
            if pd.isna(value):
                row_cells[1].text = _("N/A")
            elif isinstance(value, int | float):
                if column_name.endswith("_s"):
                    row_cells[1].text = _format_time_minutes_seconds(value)
                else:
                    row_cells[1].text = f"{value:.2f}"
            else:
                row_cells[1].text = str(value)

        # NOTE: Page break removed - was causing excessive blank space in Word reports
        # The visualizations section naturally follows the summary table
        progress_callback(2 / total_steps, _("Summary table added"))

        # Step 3: ROI Reference Map (optional) + Visualizations + ROI Event Log
        # These are extracted into helpers to keep this method concise.
        self._append_roi_reference_map(document, progress_callback, total_steps)
        self._append_visualizations(document, progress_callback, total_steps)
        self._append_roi_event_log(document, progress_callback, total_steps)
        self._append_validation_warnings(document, progress_callback, total_steps)

    def _append_roi_reference_map(
        self,
        document: DocxDocument,
        progress_callback: Callable[[float, str], None],
        total_steps: int,
    ) -> None:
        """Append ROI reference map (always shown - displays arena even without ROIs)."""
        document.add_heading(_("ROI Reference Map"), level=2)
        # Pass video_path and calibration for background frame
        fig = self.viz_generator.generate_roi_reference_plot(
            video_path=self.video_path,
            calibration=self.calibration,
        )
        memfile = io.BytesIO()
        fig.savefig(memfile, format="png", dpi=300, bbox_inches="tight")
        import matplotlib.pyplot as plt

        plt.close(fig)
        memfile.seek(0)
        # Size for 2 images per page (A4 with ~1" margins = ~6.3" usable width)
        document.add_picture(memfile, width=Inches(5.5))
        progress_callback(3 / total_steps, _("ROI map added"))

    def _append_visualizations(
        self,
        document: DocxDocument,
        progress_callback: Callable[[float, str], None],
        total_steps: int,
    ) -> None:
        """Generate and append visualization plots in parallel."""
        document.add_heading(_("Visualizations"), level=2)
        plot_configs: list[tuple[Callable[[Any], Any], str]] = [
            (
                lambda ax: self.viz_generator.generate_trajectory_plot(ax, self.video_path),
                _("Trajectory"),
            ),
            (self.viz_generator.generate_heatmap, _("Heatmap")),
            (self.viz_generator.generate_position_vs_time_plot, _("Position vs. Time")),
            (self.viz_generator.generate_cumulative_distance_plot, _("Cumulative Distance")),
            (self.viz_generator.generate_angular_velocity_plot, _("Angular Velocity")),
            (self.viz_generator.generate_thigmotaxis_plot, _("Thigmotaxis (Wall Distance)")),
        ]

        if self._should_include_geotaxis_visualization():
            plot_configs.append(
                (self.viz_generator.generate_geotaxis_plot, _("Geotaxis (Bottom Distance)"))
            )

        log.info("reporter.plots.parallel_generation.start", count=len(plot_configs))
        plot_results = self.viz_generator.generate_plots_parallel(plot_configs)

        # Names that should start on a new page (to keep 2 figures per page)
        page_break_before = {_("Cumulative Distance")}

        for i, (memfile, name) in enumerate(plot_results):
            if memfile.getbuffer().nbytes > 0:
                # Add page break before specific figures to maintain 2-per-page layout
                if name in page_break_before:
                    document.add_page_break()
                document.add_paragraph(_("Figure: {name}").format(name=name))
                # Size for 2 images per page (A4 with ~1" margins = ~6.3" usable width)
                document.add_picture(memfile, width=Inches(5.5))
            progress_callback(
                (4 + i) / total_steps,
                _("Visualization added: {name}").format(name=name),
            )

    @staticmethod
    def _normalize_aquarium_perspective(perspective: str | None) -> str:
        """Normalize perspective aliases to canonical names."""
        value = str(perspective or "").strip().lower()
        if not value:
            return "lateral"
        if value in {"top", "topdown", "top_down", "dorsal", "overhead"}:
            return "top_down"
        return "lateral"

    def _should_include_geotaxis_visualization(self) -> bool:
        """Return whether geotaxis plots should be included in DOC visualizations."""
        config = getattr(self, "behavioral_config", {}) or {}
        perspective = self._normalize_aquarium_perspective(config.get("aquarium_perspective"))
        if perspective == "top_down":
            return False
        if config.get("geotaxis_enabled") is False:
            return False
        return True

    def _append_roi_event_log(
        self,
        document: DocxDocument,
        progress_callback: Callable[[float, str], None],
        total_steps: int,
    ) -> None:
        """Append ROI event log appendix if ROI analyzer produced events."""
        if not self.r_analyzer:
            return
        document.add_page_break()
        document.add_heading(_("Appendix: ROI Event Log"), level=2)
        event_log_df = self.r_analyzer.get_event_log()
        if not event_log_df.empty:
            document.add_paragraph(
                _("Chronological log of all entries and exits from defined ROIs.")
            )

            event_log_df = event_log_df.rename(columns={"roi_name": "ROI", "event": _("Event")})

            start_time = event_log_df["timestamp"].iloc[0]

            table = document.add_table(rows=1, cols=len(event_log_df.columns))
            table.style = "Table Grid"
            for i, col_name in enumerate(event_log_df.columns):
                if col_name == "timestamp":
                    table.cell(0, i).text = _("Time (mm:ss)")
                else:
                    table.cell(0, i).text = str(col_name)

            for _index, row in event_log_df.iterrows():
                cells = table.add_row().cells
                for i, (col_name, value) in enumerate(row.items()):
                    if col_name == "timestamp":
                        elapsed = (value - start_time).total_seconds()
                        minutes = int(elapsed // 60)
                        seconds = elapsed % 60
                        cells[i].text = f"{minutes:02d}:{seconds:06.3f}"
                    else:
                        cells[i].text = str(value)
        else:
            document.add_paragraph(_("No ROI entry or exit events were recorded."))
        progress_callback(9 / total_steps, _("Event log added"))

    def _append_validation_warnings(
        self,
        document: DocxDocument,
        progress_callback: Callable[[float, str], None],
        total_steps: int,
    ) -> None:
        """Append trajectory validation warnings and technical stats."""
        document.add_page_break()
        document.add_heading(_("Appendix: Trajectory Validation"), level=2)

        # 1. Technical Stats Summary
        if self.validation_stats:
            document.add_heading(_("Quality Metrics"), level=3)
            stats = self.validation_stats

            # Create a small table for key metrics
            table = document.add_table(rows=0, cols=2)
            table.style = "Table Grid"

            def add_stat_row(label, value):
                row = table.add_row().cells
                row[0].text = str(label)
                row[1].text = str(value)

            add_stat_row(_("Total Frames Processed"), stats.get("total_frames", "N/A"))

            if "frame_range" in stats:
                fr = stats["frame_range"]
                add_stat_row(
                    _("Frame Range"),
                    f"{fr.get('min', 0)} - {fr.get('max', 0)} ({fr.get('span', 0)} frames)",
                )

            if "temporal_coverage" in stats:
                coverage = stats["temporal_coverage"] * 100
                add_stat_row(_("Temporal Coverage"), f"{coverage:.1f}%")

            if "unique_tracks" in stats:
                add_stat_row(_("Unique Track IDs"), stats["unique_tracks"])

            if "temporal_gaps" in stats:
                gaps = stats["temporal_gaps"]
                add_stat_row(
                    _("Temporal Gaps"),
                    f"{gaps.get('count', 0)} (Max: {gaps.get('max_gap_frames', 0)} frames)",
                )

                expected_gap_frames = gaps.get("expected_gap_frames")
                if expected_gap_frames is not None:
                    add_stat_row(_("Expected Gap (interval)"), f"{expected_gap_frames} frames")

                expected_count = gaps.get("expected_skip_count")
                if expected_count is not None:
                    add_stat_row(_("Expected Temporal Gaps"), str(expected_count))

                anomalous_count = gaps.get("anomalous_count")
                if anomalous_count is not None:
                    add_stat_row(_("Anomalous Temporal Gaps"), str(anomalous_count))

                anomalous_intervals = gaps.get("anomalous_intervals")
                if isinstance(anomalous_intervals, list) and anomalous_intervals:
                    document.add_paragraph(_("Anomalous gap intervals:"))
                    for interval in anomalous_intervals:
                        if not isinstance(interval, dict):
                            continue
                        from_frame = interval.get("from_frame", "?")
                        to_frame = interval.get("to_frame", "?")
                        gap_frames = interval.get("gap_frames", "?")
                        missing_frames = interval.get("missing_frames", "?")
                        document.add_paragraph(
                            f"{from_frame} → {to_frame} "
                            f"(gap={gap_frames}, missing≈{missing_frames})",
                            style="List Bullet",
                        )

            document.add_paragraph()  # Spacer

        # 2. Detailed Warnings
        document.add_heading(_("Validation Details"), level=3)
        if self.validation_warnings:
            document.add_paragraph(
                _(
                    "The following issues were detected during trajectory validation. "
                    "These may affect the precision of the calculated metrics."
                )
            )
            document.add_paragraph(
                "Definitions:\n"
                "- Temporal Gaps: Frames where the fish was not detected (occlusion/error).\n"
                "- Missing Frames: Total frames skipped if using processing_interval > 1 "
                "(expected behavior).\n"
                "- Max Gap: The largest consecutive sequence of lost frames."
            )
            for warning in self.validation_warnings:
                document.add_paragraph(str(warning), style="List Bullet")
        else:
            document.add_paragraph(
                _("No significant issues were detected during trajectory validation.")
            )

        progress_callback(10 / total_steps, _("Validation details added"))

    def export_interactive_html_report(self, output_path: Path | str) -> None:
        """
        Generate interactive HTML report with Plotly visualizations.

        IMPROVEMENT #4: Creates a standalone HTML file with interactive,
        zoomable plots that can be shared via email or web hosting.

        Features:
        - Interactive trajectory plot with velocity colormap
        - Velocity over time with freezing episodes highlighted
        - ROI occupancy bar chart
        - Hoverable tooltips with detailed information
        - Export to PNG button on each plot

        Args:
            output_path: Output file path for the HTML report
        """
        try:
            import plotly.graph_objects as go
            from plotly.subplots import make_subplots
        except ImportError:
            log.error(
                "reporter.plotly_not_installed",
                message="Plotly is required for interactive HTML reports. "
                "Install with: pip install plotly",
            )
            raise ImportError("Plotly is not installed. Run: pip install plotly") from None

        output_path = Path(output_path) if isinstance(output_path, str) else output_path

        # Ensure we have necessary data
        if self.b_analyzer is None:
            raise ValueError("BehaviorAnalyzer not available. Cannot generate HTML report.")

        trajectory_df = self.b_analyzer.trajectory_data

        # Create subplots
        fig = make_subplots(
            rows=2,
            cols=2,
            subplot_titles=(
                "Trajectory (Color = Velocity)",
                "Velocity Over Time",
                "ROI Time Spent",
                "Freezing Episodes",
            ),
            specs=[
                [{"type": "scatter"}, {"type": "scatter"}],
                [{"type": "bar"}, {"type": "scatter"}],
            ],
            vertical_spacing=0.12,
            horizontal_spacing=0.1,
        )

        # 1. Trajectory plot (top-left)
        if "x_cm" in trajectory_df.columns and "y_cm" in trajectory_df.columns:
            velocity = self.b_analyzer.calculate_velocity_timeseries()["v_mag"]
            fig.add_trace(
                go.Scatter(
                    x=trajectory_df["x_cm"],
                    y=trajectory_df["y_cm"],
                    mode="lines+markers",
                    name="Trajectory",
                    marker=dict(
                        size=4,
                        color=velocity,
                        colorscale="Viridis",
                        colorbar=dict(title="Velocity (cm/s)", x=0.46),
                        showscale=True,
                    ),
                    line=dict(width=1),
                    hovertemplate="<b>Position</b><br>X: %{x:.2f} cm<br>Y: %{y:.2f} cm<br>"
                    + "Velocity: %{marker.color:.2f} cm/s<extra></extra>",
                ),
                row=1,
                col=1,
            )
            fig.update_xaxis(title_text="X Position (cm)", row=1, col=1)
            fig.update_yaxis(title_text="Y Position (cm)", row=1, col=1)

        # 2. Velocity over time (top-right)
        velocity_ts = self.b_analyzer.calculate_velocity_timeseries()["v_mag"]
        fps = float((self.metadata or {}).get("fps", 30.0) or 30.0)
        time_seconds = trajectory_df.index.to_numpy() / fps
        fig.add_trace(
            go.Scatter(
                x=time_seconds,
                y=velocity_ts,
                mode="lines",
                name="Velocity",
                line=dict(color="blue", width=1.5),
                hovertemplate="<b>Velocity</b><br>Time: %{x:.2f} s<br>"
                + "Speed: %{y:.2f} cm/s<extra></extra>",
            ),
            row=1,
            col=2,
        )

        # Add freezing threshold line
        freezing_threshold = self.freezing_threshold or 1.5
        fig.add_hline(
            y=freezing_threshold,
            line_dash="dash",
            line_color="red",
            annotation_text=f"Freezing Threshold ({freezing_threshold} cm/s)",
            row=1,
            col=2,
        )

        fig.update_xaxis(title_text="Time (seconds)", row=1, col=2)
        fig.update_yaxis(title_text="Velocity (cm/s)", row=1, col=2)

        # 3. ROI time spent (bottom-left)
        if self.r_analyzer is not None:
            roi_time = self.r_analyzer.get_time_spent_in_rois()
            if roi_time:
                roi_names = list(roi_time.keys())
                roi_values = [roi_time[name] for name in roi_names]

                fig.add_trace(
                    go.Bar(
                        x=roi_names,
                        y=roi_values,
                        name="ROI Time",
                        marker=dict(color="lightblue"),
                        hovertemplate="<b>%{x}</b><br>Time: %{y:.2f} s<extra></extra>",
                    ),
                    row=2,
                    col=1,
                )
                fig.update_xaxis(title_text="ROI Name", row=2, col=1)
                fig.update_yaxis(title_text="Time Spent (seconds)", row=2, col=1)

        # 4. Freezing episodes (bottom-right)
        freezing_list = self.b_analyzer.detect_freezing_episodes(
            min_duration=self.freezing_duration or 1.0,
            vel_threshold=freezing_threshold,
        )
        freezing_episodes = pd.DataFrame(freezing_list)

        if not freezing_episodes.empty:
            # Plot freezing episodes as scatter points
            fps = float((self.metadata or {}).get("fps", 30.0) or 30.0)
            freeze_times = freezing_episodes["start_frame"].to_numpy() / fps
            freeze_durations = freezing_episodes["duration_s"].to_numpy()

            fig.add_trace(
                go.Scatter(
                    x=freeze_times,
                    y=freeze_durations,
                    mode="markers",
                    name="Freezing Episodes",
                    marker=dict(size=10, color="red", symbol="circle"),
                    hovertemplate="<b>Freezing Episode</b><br>Start: %{x:.2f} s<br>"
                    + "Duration: %{y:.2f} s<extra></extra>",
                ),
                row=2,
                col=2,
            )
            fig.update_xaxis(title_text="Time (seconds)", row=2, col=2)
            fig.update_yaxis(title_text="Duration (seconds)", row=2, col=2)

        # Update layout
        metadata = self.metadata or {}
        experiment_id = metadata.get("experiment_id", "Unknown")
        subject_id = metadata.get("subject_id", "Unknown")

        fig.update_layout(
            title_text=f"Interactive Analysis Report - {experiment_id} (Subject: {subject_id})",
            title_font_size=20,
            showlegend=False,
            height=900,
            width=1400,
            hovermode="closest",
        )

        # Export as standalone HTML
        html_path = str(output_path).replace(".html", "") + ".html"
        fig.write_html(
            html_path,
            config={
                "displayModeBar": True,
                "toImageButtonOptions": {
                    "format": "png",
                    "filename": f"report_{experiment_id}",
                    "height": 900,
                    "width": 1400,
                    "scale": 2,
                },
                "modeBarButtonsToAdd": ["hoverclosest", "hovercompare"],
            },
            include_plotlyjs="cdn",  # Use CDN for smaller file size
        )

        log.info(
            "reporter.interactive_html_exported",
            path=html_path,
            experiment_id=experiment_id,
        )

    @staticmethod
    def export_project_report(  # noqa: C901
        aggregated_df: pd.DataFrame,
        output_path: Path | str,
        roi_colors: dict[str, tuple[int, int, int]] | None = None,
        detector_params: dict[str, Any] | None = None,
    ):
        """Export aggregated project report with comparative analysis.

        Args:
            aggregated_df: Aggregated DataFrame with multiple experiments
            output_path: Output file path for the report
            roi_colors: Optional dict mapping ROI names to RGB color tuples
            detector_params: Optional dict with detection parameters used
        """
        output_path = Path(output_path) if isinstance(output_path, str) else output_path
        document = Document()
        document.add_heading("Aggregated Project Report", level=1)

        # Add ROI color legend if available
        if roi_colors:
            document.add_heading("ROI Color Legend", level=2)
            document.add_paragraph(
                "The following colors are used for Regions of Interest (ROIs) in this project:"
            )
            table = document.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "ROI Name"
            hdr_cells[1].text = "Color (RGB)"

            for roi_name, color in sorted(roi_colors.items()):
                row_cells = table.add_row().cells
                row_cells[0].text = roi_name
                row_cells[1].text = f"RGB({color[0]}, {color[1]}, {color[2]})"

            document.add_page_break()

        # Add detection parameters if available
        if detector_params:
            document.add_heading("Detection Parameters", level=2)
            document.add_paragraph(
                "The following parameters were used for animal detection and tracking:"
            )
            table = document.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Parameter"
            hdr_cells[1].text = "Value"

            # Sort parameters for consistent display
            for param_name, param_value in sorted(detector_params.items()):
                row_cells = table.add_row().cells
                # Format parameter name (use display mapping if available, else generic title case)
                formatted_name = DISPLAY_COLUMN_MAPPING.get(
                    param_name, param_name.replace("_", " ").title()
                )
                row_cells[0].text = formatted_name

                # Format value (handle booleans, numbers, strings)
                if isinstance(param_value, bool):
                    row_cells[1].text = "Yes" if param_value else "No"
                elif isinstance(param_value, float):
                    row_cells[1].text = f"{param_value:.2f}"
                else:
                    row_cells[1].text = str(param_value)

            document.add_page_break()

        document.add_heading("Descriptive Statistics by Group", level=2)
        # Helper list of metrics to include in the report
        metrics_of_interest = [
            "total_distance_cm",
            "mean_speed_cm_s",
            "max_speed_cm_s",
            "sharp_turns_count",
            "total_roi_entries",
            "geotaxis_bottom_zones_pct",  # Include general geotaxis if available
            "thigmotaxis_time_near_wall_pct",
        ]

        # Add dynamic Geotaxis zone columns if present
        for col in aggregated_df.columns:
            if col.startswith("geotaxis_zone_") and col.endswith("_pct"):
                metrics_of_interest.append(col)
            elif "Mean Speed in" in col or "Time in" in col:
                # Handle dynamic ROI columns if already renamed
                # (unlikely in raw but checking) or raw ROI columns
                pass

        # Filter metrics that actually exist in the dataframe
        available_metrics = [m for m in metrics_of_interest if m in aggregated_df.columns]

        # Also include any numeric columns that appear to be ROI metrics or Geotaxis
        # zones from raw data
        # (The aggregated_df here comes from Parquet so it has raw names)
        for col in aggregated_df.columns:
            if col not in available_metrics:
                if (
                    (col.startswith("time_in_") and "_s" in col)
                    or (col.startswith("entries_in_"))
                    or (col.startswith("geotaxis_zone_") and "_pct" in col)
                ):
                    if col not in available_metrics:
                        available_metrics.append(col)

        if not available_metrics:
            document.add_paragraph("No metrics available for descriptive statistics.")

        for metric in sorted(available_metrics):
            # Calculate stats
            desc_stats = aggregated_df.groupby("group_id")[metric].agg(["mean", "std", "count"])

            # Get display name
            display_name = DISPLAY_COLUMN_MAPPING.get(metric)
            if not display_name:
                # Try dynamic mapping
                display_name = DataTransformer()._translate_english_to_display(metric)

            document.add_paragraph(f"Statistics for {display_name}:", style="Heading 3")

            table = document.add_table(rows=1, cols=len(desc_stats.columns) + 1)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = DISPLAY_COLUMN_MAPPING.get("group_id", "Group")

            for i, col_name in enumerate(desc_stats.columns):
                hdr_cells[i + 1].text = col_name.capitalize()

            for index, row in desc_stats.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(index)
                for i, value in enumerate(row):
                    row_cells[i + 1].text = f"{value:.2f}"

            document.add_paragraph("")  # Spacing
        document.add_page_break()
        document.add_heading("Comparative Plots", level=2)

        metrics_for_boxplot = [
            "total_distance_cm",
            "mean_speed_cm_s",
            "max_speed_cm_s",
            "sharp_turns_count",
            "sharp_turns_per_minute",
            "total_roi_entries",
        ]

        for metric in metrics_for_boxplot:
            if metric in aggregated_df.columns:
                display_name = DISPLAY_COLUMN_MAPPING.get(metric, metric.replace("_", " ").title())
                title = f"Comparison of {display_name}"
                boxplot_fig = VisualizationGenerator.generate_comparative_boxplot(
                    aggregated_df,
                    metric,
                    title,
                )
                memfile = io.BytesIO()
                boxplot_fig.savefig(memfile, format="png", dpi=300, bbox_inches="tight")
                import matplotlib.pyplot as plt

                plt.close(boxplot_fig)
                memfile.seek(0)
                document.add_paragraph(title, style="Heading 3")
                document.add_picture(memfile, width=Inches(6.0))

        file_path = f"{output_path}.docx"
        document.save(file_path)

    @staticmethod
    def export_multi_aquarium_reports(
        results_by_aquarium: dict[int, "AnalysisResult | None"],
        output_dirs_by_aquarium: dict[int, "Path"],
        base_name: str,
        aquarium_configs: list | None = None,
        settings_obj=None,
    ) -> dict[int, dict[str, str]]:
        """
        Export separate reports for each aquarium in multi-aquarium mode.

        This method creates individual analysis reports for each aquarium,
        stored in their respective output directories.

        Args:
            results_by_aquarium: Dictionary mapping aquarium_id to AnalysisResult
                (or None if analysis failed for that aquarium).
            output_dirs_by_aquarium: Dictionary mapping aquarium_id to output Path.
            base_name: Base name for output files (e.g., video name without extension).
            aquarium_configs: Optional list of AquariumConfig objects for metadata.
            settings_obj: Optional Settings instance for Reporter configuration.

        Returns:
            Dictionary mapping aquarium_id to output paths:
            {
                0: {"summary_path": "/path/to/summary.xlsx", "report_path": "/path/to/report.docx"},
                1: {"summary_path": "/path/to/summary.xlsx", "report_path": "/path/to/report.docx"},
            }

        Example:
            >>> results = analysis_service.run_multi_aquarium_analysis(aquarium_map)
            >>> output_dirs = {0: Path("/output/aq0"), 1: Path("/output/aq1")}
            >>> paths = Reporter.export_multi_aquarium_reports(
            ...     results, output_dirs, "video_001", aquarium_configs
            ... )
        """
        from pathlib import Path

        output_paths: dict[int, dict[str, str]] = {}

        for aq_id, result in results_by_aquarium.items():
            if result is None:
                log.warning(
                    "reporter.multi_aquarium.skipping_failed",
                    aquarium_id=aq_id,
                    reason="Analysis result is None",
                )
                continue

            output_dir = output_dirs_by_aquarium.get(aq_id)
            if not output_dir:
                log.warning(
                    "reporter.multi_aquarium.no_output_dir",
                    aquarium_id=aq_id,
                )
                continue

            # Ensure output directory exists
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)

            # Get config for this aquarium if available
            config = None
            if aquarium_configs:
                config = next(
                    (c for c in aquarium_configs if getattr(c, "aquarium_id", None) == aq_id),
                    None,
                )

            # Build filename suffix with metadata
            if config:
                group = getattr(config, "group", f"aq{aq_id}")
                subject = getattr(config, "subject_id", "")
                suffix = f"_{group}_{subject}" if subject else f"_{group}"
            else:
                suffix = f"_aq{aq_id}"

            output_base = f"{base_name}{suffix}"

            try:
                # Create reporter from analysis result
                reporter = Reporter.from_analysis(result)

                # Export summary data (Excel)
                summary_path = output_dir / f"{output_base}_summary.xlsx"
                reporter.export_summary_data(str(summary_path))

                # Export individual report (Word)
                report_path = output_dir / f"{output_base}_report.docx"
                reporter.export_individual_report(str(report_path))

                output_paths[aq_id] = {
                    "summary_path": str(summary_path),
                    "report_path": str(report_path),
                }

                log.info(
                    "reporter.multi_aquarium.exported",
                    aquarium_id=aq_id,
                    summary_path=str(summary_path),
                    report_path=str(report_path),
                )

            except Exception as e:
                log.error(
                    "reporter.multi_aquarium.export_failed",
                    aquarium_id=aq_id,
                    error=str(e),
                    exc_info=True,
                )

        log.info(
            "reporter.multi_aquarium.summary",
            total_aquariums=len(results_by_aquarium),
            exported=len(output_paths),
        )

        return output_paths

    # =========================================================================
    # Phase 1.3: R/Python Export Methods
    # =========================================================================

    def export_for_r(
        self,
        output_path: Path | str,
        include_script: bool = True,
    ) -> dict[str, Path]:
        """Export data in R-friendly formats (Feather and RDS-compatible).

        Creates:
        - data.feather: Fast, efficient format readable by R's arrow::read_feather()
        - data.csv: Universal fallback
        - analysis_script.R: Template R script for loading and analyzing data

        Args:
            output_path: Directory to export files to.
            include_script: If True, includes a template R script.

        Returns:
            Dictionary with paths to created files.
        """
        import pyarrow.feather as feather

        output_dir = Path(output_path) if isinstance(output_path, str) else output_path
        output_dir.mkdir(parents=True, exist_ok=True)

        created_files: dict[str, Path] = {}

        # Export as Feather (fast, efficient for R)
        feather_path = output_dir / "data.feather"
        feather.write_feather(
            self.tidy_data,
            feather_path,
            compression="zstd",
        )
        created_files["feather"] = feather_path
        log.info("reporter.export_r.feather_saved", path=str(feather_path))

        # Export as CSV (universal fallback)
        csv_path = output_dir / "data.csv"
        self.tidy_data.to_csv(csv_path, index=False)
        created_files["csv"] = csv_path

        if include_script:
            script_path = output_dir / "analysis_script.R"
            r_script = self._generate_r_script_template()
            script_path.write_text(r_script, encoding="utf-8")
            created_files["script"] = script_path
            log.info("reporter.export_r.script_saved", path=str(script_path))

        log.info(
            "reporter.export_r.complete",
            output_dir=str(output_dir),
            files=list(created_files.keys()),
        )

        return created_files

    def export_for_python(
        self,
        output_path: Path | str,
        include_script: bool = True,
    ) -> dict[str, Path]:
        """Export data in Python-friendly formats.

        Creates:
        - data.parquet: Efficient columnar format for pandas/polars
        - data.feather: Alternative fast format
        - analysis_notebook.py: Template Python script for Jupyter/VS Code

        Args:
            output_path: Directory to export files to.
            include_script: If True, includes a template Python script.

        Returns:
            Dictionary with paths to created files.
        """
        import pyarrow.feather as feather

        output_dir = Path(output_path) if isinstance(output_path, str) else output_path
        output_dir.mkdir(parents=True, exist_ok=True)

        created_files: dict[str, Path] = {}

        # Export as Parquet (best for Python)
        parquet_path = output_dir / "data.parquet"
        self.tidy_data.to_parquet(parquet_path, index=False)
        created_files["parquet"] = parquet_path
        log.info("reporter.export_python.parquet_saved", path=str(parquet_path))

        # Export as Feather (cross-platform compatibility)
        feather_path = output_dir / "data.feather"
        feather.write_feather(self.tidy_data, feather_path, compression="zstd")
        created_files["feather"] = feather_path

        if include_script:
            script_path = output_dir / "analysis_notebook.py"
            py_script = self._generate_python_script_template()
            script_path.write_text(py_script, encoding="utf-8")
            created_files["script"] = script_path
            log.info("reporter.export_python.script_saved", path=str(script_path))

        log.info(
            "reporter.export_python.complete",
            output_dir=str(output_dir),
            files=list(created_files.keys()),
        )

        return created_files

    def _generate_r_script_template(self) -> str:
        """Generate template R script for data analysis."""
        return """# ZebTrack-AI Analysis Script for R
# Generated automatically - customize as needed

# Required packages
if (!require("arrow")) install.packages("arrow")
if (!require("ggplot2")) install.packages("ggplot2")
if (!require("dplyr")) install.packages("dplyr")

library(arrow)
library(ggplot2)
library(dplyr)

# Load data (Feather is fastest, CSV is universal fallback)
data <- arrow::read_feather("data.feather")
# Alternative: data <- read.csv("data.csv")

# Preview data structure
str(data)
summary(data)

# ============================================================
# Basic Movement Analysis
# ============================================================

# Calculate total distance per subject
distance_summary <- data %>%
  group_by(subject_id) %>%
  summarise(
    total_distance_cm = sum(velocity_cm_s * (1/30), na.rm = TRUE),  # 30 fps
    mean_velocity = mean(velocity_cm_s, na.rm = TRUE),
    max_velocity = max(velocity_cm_s, na.rm = TRUE),
    time_in_center_pct = mean(in_center_roi, na.rm = TRUE) * 100
  )

print(distance_summary)

# ============================================================
# Trajectory Plot
# ============================================================

ggplot(data, aes(x = x_cm, y = y_cm, color = as.factor(subject_id))) +
  geom_path(alpha = 0.5) +
  labs(
    title = "Zebrafish Trajectories",
    x = "X Position (cm)",
    y = "Y Position (cm)",
    color = "Subject"
  ) +
  theme_minimal()

# ============================================================
# Velocity Over Time
# ============================================================

ggplot(data, aes(x = timestamp, y = velocity_cm_s)) +
  geom_line(alpha = 0.5) +
  geom_smooth(method = "loess", se = TRUE) +
  labs(
    title = "Velocity Over Time",
    x = "Time (seconds)",
    y = "Velocity (cm/s)"
  ) +
  theme_minimal()

# ============================================================
# Statistical Tests (Example)
# ============================================================

# If you have experimental groups, add group column and run tests
# Example:
# t.test(velocity_cm_s ~ group, data = data)
# wilcox.test(velocity_cm_s ~ group, data = data)
"""

    def _generate_python_script_template(self) -> str:
        """Generate template Python script for data analysis."""
        return """# %% [markdown]
# # ZebTrack-AI Analysis Notebook
# Generated automatically - customize as needed

# %% [markdown]
# ## Load Data

# %%
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns

# Load data (Parquet is fastest)
df = pd.read_parquet("data.parquet")
# Alternative: df = pd.read_feather("data.feather")

# Preview data
print(f"Shape: {df.shape}")
print(f"Columns: {list(df.columns)}")
df.head()

# %%
df.describe()

# %% [markdown]
# ## Basic Movement Analysis

# %%
# Calculate summary statistics per subject
summary = df.groupby("subject_id").agg({
    "velocity_cm_s": ["mean", "max", "std"],
    "x_cm": ["min", "max"],
    "y_cm": ["min", "max"],
}).round(2)

summary.columns = ["_".join(col) for col in summary.columns]
print(summary)

# %% [markdown]
# ## Trajectory Visualization

# %%
fig, ax = plt.subplots(figsize=(10, 8))

for subject_id in df["subject_id"].unique():
    subject_data = df[df["subject_id"] == subject_id]
    ax.plot(
        subject_data["x_cm"],
        subject_data["y_cm"],
        alpha=0.6,
        label=f"Subject {subject_id}"
    )

ax.set_xlabel("X Position (cm)")
ax.set_ylabel("Y Position (cm)")
ax.set_title("Zebrafish Trajectories")
ax.legend()
ax.set_aspect("equal")
plt.tight_layout()
plt.show()

# %% [markdown]
# ## Velocity Analysis

# %%
fig, axes = plt.subplots(1, 2, figsize=(14, 5))

# Velocity over time
ax1 = axes[0]
ax1.plot(df["timestamp"], df["velocity_cm_s"], alpha=0.3, linewidth=0.5)
# Rolling average
rolling = df["velocity_cm_s"].rolling(window=30, min_periods=1).mean()
ax1.plot(df["timestamp"], rolling, color="red", linewidth=2, label="30-frame avg")
ax1.set_xlabel("Time (s)")
ax1.set_ylabel("Velocity (cm/s)")
ax1.set_title("Velocity Over Time")
ax1.legend()

# Velocity distribution
ax2 = axes[1]
ax2.hist(df["velocity_cm_s"].dropna(), bins=50, edgecolor="black", alpha=0.7)
ax2.axvline(df["velocity_cm_s"].mean(), color="red", linestyle="--", label="Mean")
ax2.set_xlabel("Velocity (cm/s)")
ax2.set_ylabel("Frequency")
ax2.set_title("Velocity Distribution")
ax2.legend()

plt.tight_layout()
plt.show()

# %% [markdown]
# ## Export Results

# %%
# Save summary to CSV for further analysis
summary.to_csv("analysis_summary.csv")
print("Summary saved to analysis_summary.csv")

# %% [markdown]
# ## Statistical Tests (Add Your Groups)

# %%
# If you have experimental groups, uncomment and modify:
# from scipy import stats
#
# group_a = df[df["group"] == "control"]["velocity_cm_s"]
# group_b = df[df["group"] == "treatment"]["velocity_cm_s"]
#
# t_stat, p_value = stats.ttest_ind(group_a, group_b)
# print(f"T-test: t={t_stat:.3f}, p={p_value:.4f}")
"""

    # =========================================================================
    # API Aliases for Documentation Consistency
    # =========================================================================

    def export_feather(self, output_path: Path | str) -> Path:
        """Export data as Feather format for fast R/Python loading.

        Convenience alias that wraps export_for_r() to export only the Feather file.
        Mentioned in CHANGELOG.md Phase 1 as a new export format.

        Args:
            output_path: Path to the output .feather file.

        Returns:
            Path to the created Feather file.

        Example:
            >>> reporter.export_feather("output/data.feather")
        """
        import pyarrow.feather as feather

        output_file = Path(output_path) if isinstance(output_path, str) else output_path
        output_file.parent.mkdir(parents=True, exist_ok=True)

        feather.write_feather(self.tidy_data, output_file, compression="zstd")
        log.info("reporter.export_feather.saved", path=str(output_file))

        return output_file

    def export_r_script(self, output_path: Path | str) -> Path:
        """Export template R script for statistical analysis.

        Convenience alias that extracts only the R script generation.
        Mentioned in CHANGELOG.md Phase 1 as a new export format.

        Args:
            output_path: Path to the output .R file.

        Returns:
            Path to the created R script file.

        Example:
            >>> reporter.export_r_script("output/analysis.R")
        """
        output_file = Path(output_path) if isinstance(output_path, str) else output_path
        output_file.parent.mkdir(parents=True, exist_ok=True)

        r_script = self._generate_r_script_template()
        output_file.write_text(r_script, encoding="utf-8")
        log.info("reporter.export_r_script.saved", path=str(output_file))

        return output_file

    def export_python_script(self, output_path: Path | str) -> Path:
        """Export template Python script for Jupyter/VS Code analysis.

        Convenience alias that extracts only the Python script generation.
        Mentioned in CHANGELOG.md Phase 1 as a new export format.

        Args:
            output_path: Path to the output .py file.

        Returns:
            Path to the created Python script file.

        Example:
            >>> reporter.export_python_script("output/analysis_notebook.py")
        """
        output_file = Path(output_path) if isinstance(output_path, str) else output_path
        output_file.parent.mkdir(parents=True, exist_ok=True)

        py_script = self._generate_python_script_template()
        output_file.write_text(py_script, encoding="utf-8")
        log.info("reporter.export_python_script.saved", path=str(output_file))

        return output_file
