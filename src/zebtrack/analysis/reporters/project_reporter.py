"""Project-level report functions.

These were formerly ``@staticmethod`` members of ``Reporter``.  They
operate on aggregated DataFrames or multi-aquarium result dicts and
do **not** require a ``ReporterContext`` instance.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any

import pandas as pd
import structlog
from docx import Document
from docx.shared import Inches

from zebtrack.analysis.data_transformer import DISPLAY_COLUMN_MAPPING, DataTransformer
from zebtrack.analysis.visualization_generator import VisualizationGenerator

log = structlog.get_logger(__name__)


def export_project_report(
    aggregated_df: pd.DataFrame,
    output_path: Path | str,
    roi_colors: dict[str, tuple[int, int, int]] | None = None,
    detector_params: dict[str, Any] | None = None,
) -> None:
    """Export aggregated project report with comparative analysis.

    Args:
        aggregated_df: Aggregated DataFrame with multiple experiments.
        output_path: Output file path for the report.
        roi_colors: Optional dict mapping ROI names to RGB colour tuples.
        detector_params: Optional dict with detection parameters used.
    """
    output_path = Path(output_path) if isinstance(output_path, str) else output_path
    document = Document()
    document.add_heading("Aggregated Project Report", level=1)

    # ROI colour legend
    if roi_colors:
        document.add_heading("ROI Color Legend", level=2)
        document.add_paragraph(
            "The following colors are used for Regions of Interest (ROIs) in this project:"
        )
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "ROI Name"
        hdr_cells[1].text = "Color (RGB)"

        for roi_name, color in sorted(roi_colors.items()):
            row_cells = table.add_row().cells
            row_cells[0].text = roi_name
            row_cells[1].text = f"RGB({color[0]}, {color[1]}, {color[2]})"

        document.add_page_break()

    # Detection parameters
    if detector_params:
        document.add_heading("Detection Parameters", level=2)
        document.add_paragraph(
            "The following parameters were used for animal detection and tracking:"
        )
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Parameter"
        hdr_cells[1].text = "Value"

        for param_name, param_value in sorted(detector_params.items()):
            row_cells = table.add_row().cells
            formatted_name = DISPLAY_COLUMN_MAPPING.get(
                param_name, param_name.replace("_", " ").title()
            )
            row_cells[0].text = formatted_name

            if isinstance(param_value, bool):
                row_cells[1].text = "Yes" if param_value else "No"
            elif isinstance(param_value, float):
                row_cells[1].text = f"{param_value:.2f}"
            else:
                row_cells[1].text = str(param_value)

        document.add_page_break()

    # Descriptive statistics by group
    document.add_heading("Descriptive Statistics by Group", level=2)

    metrics_of_interest = [
        "total_distance_cm",
        "mean_speed_cm_s",
        "max_speed_cm_s",
        "sharp_turns_count",
        "total_roi_entries",
        "geotaxis_bottom_zones_pct",
        "thigmotaxis_time_near_wall_pct",
    ]

    for col in aggregated_df.columns:
        if col.startswith("geotaxis_zone_") and col.endswith("_pct"):
            metrics_of_interest.append(col)

    available_metrics = [m for m in metrics_of_interest if m in aggregated_df.columns]

    for col in aggregated_df.columns:
        if col not in available_metrics:
            if (
                (col.startswith("time_in_") and "_s" in col)
                or (col.startswith("entries_in_"))
                or (col.startswith("geotaxis_zone_") and "_pct" in col)
            ):
                available_metrics.append(col)

    if not available_metrics:
        document.add_paragraph("No metrics available for descriptive statistics.")

    for metric in sorted(available_metrics):
        desc_stats = aggregated_df.groupby("group_id")[metric].agg(["mean", "std", "count"])

        display_name = DISPLAY_COLUMN_MAPPING.get(metric)
        if not display_name:
            display_name = DataTransformer()._translate_english_to_display(metric)

        document.add_paragraph(f"Statistics for {display_name}:", style="Heading 3")

        table = document.add_table(rows=1, cols=len(desc_stats.columns) + 1)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = DISPLAY_COLUMN_MAPPING.get("group_id", "Group")

        for i, col_name in enumerate(desc_stats.columns):
            hdr_cells[i + 1].text = col_name.capitalize()

        for index, row in desc_stats.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(index)
            for i, value in enumerate(row):
                row_cells[i + 1].text = f"{value:.2f}"

        document.add_paragraph("")

    document.add_page_break()
    document.add_heading("Comparative Plots", level=2)

    metrics_for_boxplot = [
        "total_distance_cm",
        "mean_speed_cm_s",
        "max_speed_cm_s",
        "sharp_turns_count",
        "sharp_turns_per_minute",
        "total_roi_entries",
    ]

    for metric in metrics_for_boxplot:
        if metric in aggregated_df.columns:
            display_name = DISPLAY_COLUMN_MAPPING.get(metric, metric.replace("_", " ").title())
            title = f"Comparison of {display_name}"
            boxplot_fig = VisualizationGenerator.generate_comparative_boxplot(
                aggregated_df,
                metric,
                title,
            )
            memfile = io.BytesIO()
            boxplot_fig.savefig(memfile, format="png", dpi=300, bbox_inches="tight")
            import matplotlib.pyplot as plt

            plt.close(boxplot_fig)
            memfile.seek(0)
            document.add_paragraph(title, style="Heading 3")
            document.add_picture(memfile, width=Inches(6.0))

    file_path = f"{output_path}.docx"
    document.save(file_path)


def export_multi_aquarium_reports(
    results_by_aquarium: dict[int, Any | None],
    output_dirs_by_aquarium: dict[int, Path],
    base_name: str,
    aquarium_configs: list | None = None,
    settings_obj: Any = None,
) -> dict[int, dict[str, str]]:
    """Export separate reports for each aquarium in multi-aquarium mode.

    Args:
        results_by_aquarium: ``{aquarium_id: AnalysisResult | None}``.
        output_dirs_by_aquarium: ``{aquarium_id: output_dir_path}``.
        base_name: Base name for output files (e.g. video stem).
        aquarium_configs: Optional list of ``AquariumConfig`` objects.
        settings_obj: Optional Settings instance.

    Returns:
        ``{aquarium_id: {"summary_path": ..., "report_path": ...}}``.
    """
    from zebtrack.analysis.reporters.excel_reporter import ExcelReporter
    from zebtrack.analysis.reporters.reporter_context import ReporterContext
    from zebtrack.analysis.reporters.word_reporter import WordReporter

    output_paths: dict[int, dict[str, str]] = {}

    for aq_id, result in results_by_aquarium.items():
        if result is None:
            log.warning(
                "reporter.multi_aquarium.skipping_failed",
                aquarium_id=aq_id,
                reason="Analysis result is None",
            )
            continue

        output_dir = output_dirs_by_aquarium.get(aq_id)
        if not output_dir:
            log.warning(
                "reporter.multi_aquarium.no_output_dir",
                aquarium_id=aq_id,
            )
            continue

        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        config = None
        if aquarium_configs:
            config = next(
                (c for c in aquarium_configs if getattr(c, "aquarium_id", None) == aq_id),
                None,
            )

        if config:
            group = getattr(config, "group", f"aq{aq_id}")
            subject = getattr(config, "subject_id", "")
            suffix = f"_{group}_{subject}" if subject else f"_{group}"
        else:
            suffix = f"_aq{aq_id}"

        output_base = f"{base_name}{suffix}"

        try:
            ctx = ReporterContext.from_analysis(result)

            summary_path = output_dir / f"{output_base}_summary.xlsx"
            ExcelReporter(ctx).export_summary(str(summary_path))

            report_path = output_dir / f"{output_base}_report.docx"
            WordReporter(ctx).export_individual_report(str(report_path))

            output_paths[aq_id] = {
                "summary_path": str(summary_path),
                "report_path": str(report_path),
            }

            log.info(
                "reporter.multi_aquarium.exported",
                aquarium_id=aq_id,
                summary_path=str(summary_path),
                report_path=str(report_path),
            )

        except Exception as e:
            log.error(
                "reporter.multi_aquarium.export_failed",
                aquarium_id=aq_id,
                error=str(e),
                exc_info=True,
            )

    log.info(
        "reporter.multi_aquarium.summary",
        total_aquariums=len(results_by_aquarium),
        exported=len(output_paths),
    )

    return output_paths
