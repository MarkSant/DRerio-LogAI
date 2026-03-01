"""Word document report generation.

Generates individual experiment reports and project-level comparative
reports in ``.docx`` format using ``python-docx`` and ``docxtpl``.
"""

from __future__ import annotations

import io
import tempfile
from collections.abc import Callable
from pathlib import Path
from typing import TYPE_CHECKING, Any

import pandas as pd
import structlog
from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Inches
from docxtpl import DocxTemplate

from zebtrack.analysis.data_transformer import DISPLAY_COLUMN_MAPPING
from zebtrack.analysis.reporters.reporter_context import (
    INDIVIDUAL_REPORT_TEMPLATE,
    _,
)

if TYPE_CHECKING:
    from zebtrack.analysis.reporters.reporter_context import ReporterContext

log = structlog.get_logger(__name__)


def _format_time_minutes_seconds(seconds: float | None) -> str:
    """Format time as MM:SS or HH:MM:SS if over an hour."""
    if pd.isna(seconds) or seconds is None:
        return "N/A"

    seconds = float(seconds)
    if seconds < 60:
        return f"{int(seconds)}s"
    elif seconds < 3600:
        minutes = int(seconds // 60)
        secs = int(seconds % 60)
        return f"{minutes}:{secs:02d}"
    else:
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        return f"{hours}:{minutes:02d}:{secs:02d}"


class WordReporter:
    """Generate individual Word document reports from a ``ReporterContext``.

    Example:
        >>> ctx = ReporterContext.from_analysis(analysis_result)
        >>> WordReporter(ctx).export_individual_report("report.docx")
    """

    def __init__(self, ctx: ReporterContext) -> None:
        self._ctx = ctx

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------
    def export_individual_report(self, output_path: Path | str) -> None:
        """Export a complete individual report.

        Convenience wrapper around ``export_individual_report_step_by_step``
        for consumers that don't need progress updates.

        Args:
            output_path: Output file path for the report.
        """
        self.export_individual_report_step_by_step(output_path, lambda p, s: None)

    def export_individual_report_step_by_step(
        self, output_path: Path | str, progress_callback: Callable[[float, str], None]
    ) -> None:
        """Export individual report with step-by-step progress reporting.

        Args:
            output_path: Output file path for the report.
            progress_callback: ``(progress: float, status: str) -> None``.
        """
        output_path = Path(output_path) if isinstance(output_path, str) else output_path
        total_steps = 11
        template_path = INDIVIDUAL_REPORT_TEMPLATE
        heading_text = _("Analysis Report - {experiment_id}").format(
            experiment_id=(self._ctx.metadata or {}).get("experiment_id", "Unknown")
        )

        doc_template, document = self._prepare_report_document(template_path, heading_text)

        # Step 1: Metadata section
        self._append_metadata_section(document, progress_callback, total_steps)

        # Save
        try:
            file_path = (
                f"{output_path}.docx"
                if not str(output_path).lower().endswith(".docx")
                else str(output_path)
            )
            if doc_template is not None:
                doc_template.save(file_path)
            else:
                document.save(file_path)
            progress_callback(1.0, _("Report saved"))
            log.info("reporter.individual_report.saved", path=file_path)
        except Exception as e:
            log.error("reporter.individual_report.save_failed", error=str(e), exc_info=True)
            raise

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------
    @staticmethod
    def _prepare_report_document(
        template_path: Path, heading_text: str
    ) -> tuple[DocxTemplate | None, DocxDocument]:
        """Prepare a docx document using template if available."""
        doc_template: DocxTemplate | None = None
        document: DocxDocument | None = None
        template_rendered = False

        if template_path.exists():
            try:
                doc_template = DocxTemplate(str(template_path))
                doc_template.render({"title": heading_text})
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
                    tmp_path = Path(tmp_file.name)

                try:
                    doc_template.save(str(tmp_path))
                    document = Document(str(tmp_path))
                    template_rendered = True
                finally:
                    tmp_path.unlink(missing_ok=True)

                doc_template = None
            except Exception:
                log.warning(
                    "analysis.reporter.template_render_failed",
                    template=str(template_path),
                    exc_info=True,
                )
                document = Document()
        else:
            log.warning(
                "analysis.reporter.template_missing_fallback",
                template=str(template_path),
            )
            document = Document()

        if document is None:
            document = Document()

        if doc_template is None and not template_rendered:
            document.add_heading(heading_text, level=1)

        return doc_template, document

    def _append_metadata_section(
        self,
        document: DocxDocument,
        progress_callback: Callable[[float, str], None],
        total_steps: int,
    ) -> None:
        """Append metadata, summary table, visualisations and appendices."""
        ctx = self._ctx

        document.add_heading(_("Experiment Metadata"), level=2)
        for key, value in (ctx.metadata or {}).items():
            document.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
        progress_callback(1 / total_steps, _("Metadata added"))

        # Summary table
        document.add_heading(_("Metrics Summary"), level=2)
        df = ctx.tidy_data.drop(
            columns=[k for k in (ctx.metadata or {}).keys() if k in ctx.tidy_data.columns]
        )

        # Geotaxis column renaming for Word
        height_cm = float((ctx.metadata or {}).get("aquarium_height_cm", 0) or 0)
        num_zones = int((ctx.metadata or {}).get("geotaxis_num_zones", 0) or 0)
        if hasattr(ctx, "behavioral_config") and num_zones == 0:
            num_zones = int(ctx.behavioral_config.get("geotaxis_num_zones", 0) or 0)

        if height_cm > 0 and num_zones > 0:
            df = ctx.data_transformer.rename_geotaxis_columns(df, height_cm, num_zones)
        else:
            rename_geo: dict[str, str] = {}
            for col in df.columns:
                if col.startswith("geotaxis_zone_") and col.endswith("_pct"):
                    try:
                        idx = int(col.split("_")[2])
                        if idx == 0:
                            rename_geo[col] = "Geotaxis Zona 1 - Fundo (%)"
                        else:
                            rename_geo[col] = f"Geotaxis Zona {idx + 1} (%)"
                    except (IndexError, ValueError):
                        log.debug(
                            "reporter.geotaxis_rename.parse_error",
                            column=col,
                            exc_info=True,
                        )
            if rename_geo:
                df = df.rename(columns=rename_geo)

        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.cell(0, 0).text = _("Metric")
        table.cell(0, 1).text = _("Value")

        for column_name in df.columns:
            row_cells = table.add_row().cells
            formatted_name = DISPLAY_COLUMN_MAPPING.get(
                column_name, column_name.replace("_", " ").title()
            )
            row_cells[0].text = formatted_name

            value = df[column_name].iloc[0]
            if pd.isna(value):
                row_cells[1].text = _("N/A")
            elif isinstance(value, int | float):
                if column_name.endswith("_s"):
                    row_cells[1].text = _format_time_minutes_seconds(value)
                elif isinstance(value, int) or (
                    isinstance(value, float) and value == int(value) and "count" in column_name
                ):
                    row_cells[1].text = str(int(value))
                else:
                    row_cells[1].text = f"{value:.2f}"
            else:
                row_cells[1].text = str(value)

        progress_callback(2 / total_steps, _("Summary table added"))

        # Delegate to sub-sections
        self._append_roi_reference_map(document, progress_callback, total_steps)
        self._append_visualizations(document, progress_callback, total_steps)
        self._append_roi_event_log(document, progress_callback, total_steps)
        self._append_validation_warnings(document, progress_callback, total_steps)

    def _append_roi_reference_map(
        self,
        document: DocxDocument,
        progress_callback: Callable[[float, str], None],
        total_steps: int,
    ) -> None:
        """Append ROI reference map (always shown)."""
        document.add_heading(_("ROI Reference Map"), level=2)
        fig = self._ctx.viz_generator.generate_roi_reference_plot(
            video_path=self._ctx.video_path,
            calibration=self._ctx.calibration,
        )
        memfile = io.BytesIO()
        fig.savefig(memfile, format="png", dpi=300, bbox_inches="tight")
        import matplotlib.pyplot as plt

        plt.close(fig)
        memfile.seek(0)
        document.add_picture(memfile, width=Inches(5.5))
        progress_callback(3 / total_steps, _("ROI map added"))

    def _append_visualizations(
        self,
        document: DocxDocument,
        progress_callback: Callable[[float, str], None],
        total_steps: int,
    ) -> None:
        """Generate and append visualisation plots in parallel."""
        document.add_heading(_("Visualizations"), level=2)
        plot_configs: list[tuple[Callable[[Any], Any], str]] = [
            (
                lambda ax: self._ctx.viz_generator.generate_trajectory_plot(
                    ax, self._ctx.video_path
                ),
                _("Trajectory"),
            ),
            (self._ctx.viz_generator.generate_heatmap, _("Heatmap")),
            (self._ctx.viz_generator.generate_position_vs_time_plot, _("Position vs. Time")),
            (
                self._ctx.viz_generator.generate_cumulative_distance_plot,
                _("Cumulative Distance"),
            ),
            (
                self._ctx.viz_generator.generate_angular_velocity_plot,
                _("Angular Velocity"),
            ),
            (
                self._ctx.viz_generator.generate_thigmotaxis_plot,
                _("Thigmotaxis (Wall Distance)"),
            ),
        ]

        if self._should_include_geotaxis_visualization():
            plot_configs.append(
                (
                    self._ctx.viz_generator.generate_geotaxis_plot,
                    _("Geotaxis (Bottom Distance)"),
                )
            )

        log.info("reporter.plots.parallel_generation.start", count=len(plot_configs))
        plot_results = self._ctx.viz_generator.generate_plots_parallel(plot_configs)

        page_break_before = {_("Cumulative Distance")}

        for i, (memfile, name) in enumerate(plot_results):
            if memfile.getbuffer().nbytes > 0:
                if name in page_break_before:
                    document.add_page_break()
                document.add_paragraph(_("Figure: {name}").format(name=name))
                document.add_picture(memfile, width=Inches(5.5))
            progress_callback(
                (4 + i) / total_steps,
                _("Visualization added: {name}").format(name=name),
            )

    def _should_include_geotaxis_visualization(self) -> bool:
        """Return whether geotaxis plots should be included in DOC visualisations."""
        config = getattr(self._ctx, "behavioral_config", {}) or {}
        perspective = self._ctx._normalize_aquarium_perspective(config.get("aquarium_perspective"))
        if perspective == "top_down":
            return False
        if config.get("geotaxis_enabled") is False:
            return False
        return True

    def _append_roi_event_log(
        self,
        document: DocxDocument,
        progress_callback: Callable[[float, str], None],
        total_steps: int,
    ) -> None:
        """Append ROI event log appendix if ROI analyser produced events."""
        if not self._ctx.r_analyzer:
            return
        document.add_page_break()
        document.add_heading(_("Appendix: ROI Event Log"), level=2)
        event_log_df = self._ctx.r_analyzer.get_event_log()
        if not event_log_df.empty:
            document.add_paragraph(
                _("Chronological log of all entries and exits from defined ROIs.")
            )

            event_log_df = event_log_df.rename(columns={"roi_name": "ROI", "event": _("Event")})

            start_time = event_log_df["timestamp"].iloc[0]

            table = document.add_table(rows=1, cols=len(event_log_df.columns))
            table.style = "Table Grid"
            for i, col_name in enumerate(event_log_df.columns):
                if col_name == "timestamp":
                    table.cell(0, i).text = _("Time (mm:ss)")
                else:
                    table.cell(0, i).text = str(col_name)

            for _index, row in event_log_df.iterrows():
                cells = table.add_row().cells
                for i, (col_name, value) in enumerate(row.items()):
                    if col_name == "timestamp":
                        elapsed = (value - start_time).total_seconds()
                        minutes = int(elapsed // 60)
                        seconds = elapsed % 60
                        cells[i].text = f"{minutes:02d}:{seconds:06.3f}"
                    else:
                        cells[i].text = str(value)
        else:
            document.add_paragraph(_("No ROI entry or exit events were recorded."))
        progress_callback(9 / total_steps, _("Event log added"))

    def _append_validation_warnings(
        self,
        document: DocxDocument,
        progress_callback: Callable[[float, str], None],
        total_steps: int,
    ) -> None:
        """Append trajectory validation warnings and technical stats."""
        document.add_page_break()
        document.add_heading(_("Appendix: Trajectory Validation"), level=2)

        # 1. Technical Stats Summary
        if self._ctx.validation_stats:
            document.add_heading(_("Quality Metrics"), level=3)
            stats = self._ctx.validation_stats

            table = document.add_table(rows=0, cols=2)
            table.style = "Table Grid"

            def add_stat_row(label: str, value: object) -> None:
                row = table.add_row().cells
                row[0].text = str(label)
                row[1].text = str(value)

            add_stat_row(_("Total Frames Processed"), stats.get("total_frames", "N/A"))

            if "frame_range" in stats:
                fr = stats["frame_range"]
                add_stat_row(
                    _("Frame Range"),
                    f"{fr.get('min', 0)} - {fr.get('max', 0)} ({fr.get('span', 0)} frames)",
                )

            if "temporal_coverage" in stats:
                coverage = stats["temporal_coverage"] * 100
                add_stat_row(_("Temporal Coverage"), f"{coverage:.1f}%")

            if "unique_tracks" in stats:
                add_stat_row(_("Unique Track IDs"), stats["unique_tracks"])

            if "temporal_gaps" in stats:
                gaps = stats["temporal_gaps"]
                add_stat_row(
                    _("Temporal Gaps"),
                    f"{gaps.get('count', 0)} (Max: {gaps.get('max_gap_frames', 0)} frames)",
                )

                expected_gap_frames = gaps.get("expected_gap_frames")
                if expected_gap_frames is not None:
                    add_stat_row(_("Expected Gap (interval)"), f"{expected_gap_frames} frames")

                expected_count = gaps.get("expected_skip_count")
                if expected_count is not None:
                    add_stat_row(_("Expected Temporal Gaps"), str(expected_count))

                anomalous_count = gaps.get("anomalous_count")
                if anomalous_count is not None:
                    add_stat_row(_("Anomalous Temporal Gaps"), str(anomalous_count))

                anomalous_intervals = gaps.get("anomalous_intervals")
                if isinstance(anomalous_intervals, list) and anomalous_intervals:
                    document.add_paragraph(_("Anomalous gap intervals:"))
                    for interval in anomalous_intervals:
                        if not isinstance(interval, dict):
                            continue
                        from_frame = interval.get("from_frame", "?")
                        to_frame = interval.get("to_frame", "?")
                        gap_frames = interval.get("gap_frames", "?")
                        missing_frames = interval.get("missing_frames", "?")
                        document.add_paragraph(
                            f"{from_frame} → {to_frame} "
                            f"(gap={gap_frames}, missing≈{missing_frames})",
                            style="List Bullet",
                        )

            document.add_paragraph()  # Spacer

        # 2. Detailed Warnings
        document.add_heading(_("Validation Details"), level=3)
        if self._ctx.validation_warnings:
            document.add_paragraph(
                _(
                    "The following issues were detected during trajectory validation. "
                    "These may affect the precision of the calculated metrics."
                )
            )
            document.add_paragraph(
                "Definitions:\n"
                "- Temporal Gaps: Frames where the fish was not detected (occlusion/error).\n"
                "- Missing Frames: Total frames skipped if using processing_interval > 1 "
                "(expected behavior).\n"
                "- Max Gap: The largest consecutive sequence of lost frames."
            )
            for warning in self._ctx.validation_warnings:
                document.add_paragraph(str(warning), style="List Bullet")
        else:
            document.add_paragraph(
                _("No significant issues were detected during trajectory validation.")
            )

        progress_callback(10 / total_steps, _("Validation details added"))
