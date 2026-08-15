"""Block detail dialog for Day x Group session management (v2.3.0).

Shows all subjects (cobaias) in the block with status and quick actions.

Version: 2.3.1
"""

from __future__ import annotations

import re
import threading
from datetime import datetime
from pathlib import Path
from tkinter import Button, Canvas, Frame, Label, Toplevel, messagebox, simpledialog, ttk
from typing import TYPE_CHECKING

import structlog

from zebtrack.core.services.session_duration_resolver import (
    SUBJECT_WILDCARD,
    collect_block_durations,
    resolve_session_duration,
    set_duration_override,
)
from zebtrack.i18n import _
from zebtrack.utils.report_files import find_summary_excel_file, has_summary_excel_output

if TYPE_CHECKING:
    from zebtrack.coordinators.live_batch_coordinator import LiveBatchCoordinator
    from zebtrack.coordinators.live_camera_session_coordinator import LiveCameraSessionCoordinator

log = structlog.get_logger(__name__)


def _block_label(day_num, group_name) -> str:
    """Human-readable name of a Day x Group block.

    DISPLAY ONLY. The persisted key for the same block is built separately as
    ``Dia_{n}_{grupo}`` (see add_note) and must stay Portuguese -- it is a
    dictionary key inside project_data, not copy.
    """
    return _("Day {day} - {group}").format(day=day_num, group=group_name)


class BlockDetailDialog(Toplevel):
    """Detail dialog for Day x Group block."""

    def __init__(
        self,
        parent,
        day: int | str,
        group: str,
        project_manager,
        session_coordinator: LiveCameraSessionCoordinator,
        live_batch_coordinator: LiveBatchCoordinator,
    ):
        """Initialize block detail dialog.

        Args:
            parent: Parent widget
            day: Day number (int) or label (str, e.g., "Dia_1")
            group: Group name (e.g., "Controle")
            project_manager: ProjectManager instance for project data access
            session_coordinator: LiveCameraSessionCoordinator for session management
            live_batch_coordinator: LiveBatchCoordinator for batch tracking
        """
        super().__init__(parent)
        # v2.3.1: Handle both int and str day formats
        self.day_num = (
            day if isinstance(day, int) else int(day.replace("Dia_", "").replace("D", ""))
        )
        self.day = f"Dia_{self.day_num}" if isinstance(day, int) else str(day)
        self.group_name = group
        self.project_manager = project_manager
        self.session_coordinator = session_coordinator
        self.live_batch_coordinator = live_batch_coordinator

        # Extract experiment data from project_manager
        project_data = (
            project_manager.project_data if hasattr(project_manager, "project_data") else {}
        )
        self.subjects_per_group = project_data.get("subjects_per_group", 0)
        self.completed_sessions = (
            set(project_manager.get_completed_sessions())
            if hasattr(project_manager, "get_completed_sessions")
            else set()
        )

        # Cache project polygon status for the header indicator + per-subject
        # "reused" badge. Read once at dialog init so per-row rendering doesn't
        # re-walk the zone data structures.
        try:
            zone_data = (
                project_manager.get_zone_data()
                if hasattr(project_manager, "get_zone_data")
                else None
            )
            self._project_has_polygon = bool(zone_data and getattr(zone_data, "polygon", None))
        except Exception:
            self._project_has_polygon = False

        # v2.3.1: Debug log for session detection
        log.info(
            "block_detail.init",
            day=self.day_num,
            group=self.group_name,
            subjects_per_group=self.subjects_per_group,
            completed_sessions=list(self.completed_sessions),
            project_has_polygon=self._project_has_polygon,
            project_path=str(project_manager.project_path)
            if project_manager.project_path
            else None,
        )

        # Window config
        self.title(_("Sessions: {block}").format(block=_block_label(self.day_num, group)))
        self.geometry("700x640")
        self.transient(parent)
        self.grab_set()

        # Per-block camera override (None = use project default).
        self._camera_index_override: int | None = None
        self._camera_friendly_name_override: str | None = None
        self._camera_label: Label | None = None

        # Duration widgets (populated by build_ui). Unlike the camera override,
        # durations are PERSISTED in project_data — an experiment where animals
        # were recorded for different lengths must stay auditable after restart.
        self._duration_label: Label | None = None
        self._subject_container: Frame | None = None

        self.build_ui()

    def build_ui(self):
        """Build dialog UI."""
        # Header
        header = Frame(self, bg="#f8f9fa", height=80)
        header.pack(fill="x")
        header.pack_propagate(False)

        Label(
            header,
            text=f"📋 {_block_label(self.day_num, self.group_name)}",
            font=("Segoe UI", 14, "bold"),
            bg="#f8f9fa",
        ).pack(side="left", padx=20, pady=20)

        # Progress info - v2.3.1: Use subjects_per_group and completed_sessions
        subjects = self._subjects()
        completed = sum(
            1 for s in subjects if (self.day_num, self.group_name, s) in self.completed_sessions
        )

        Label(
            header,
            text=_("📊 Progress: {done}/{total} sessions").format(
                done=completed, total=len(subjects)
            ),
            font=("Segoe UI", 11),
            bg="#f8f9fa",
            fg="#555",
        ).pack(side="left", padx=10, pady=20)

        # Project-level polygon indicator: shows whether the project already has
        # an arena polygon defined (which gets reused across subjects in live
        # projects). Helps users know whether the first session of the block
        # will trigger zone calibration or jump straight to recording.
        polygon_text = (
            _("🏟️ Project polygon: ✅ Defined")
            if self._project_has_polygon
            else _("🏟️ Project polygon: ⚠️ Not defined")
        )
        polygon_color = "#0a7" if self._project_has_polygon else "#a23"
        Label(
            header,
            text=polygon_text,
            font=("Segoe UI", 10),
            bg="#f8f9fa",
            fg=polygon_color,
        ).pack(side="right", padx=20, pady=20)

        # Subject list
        list_frame = Frame(self)
        list_frame.pack(fill="both", expand=True, padx=20, pady=10)

        Label(
            list_frame,
            text=_("🐟 Subjects"),
            font=("Segoe UI", 11, "bold"),
        ).pack(anchor="w", pady=(0, 10))

        # Canvas + Scrollbar
        canvas = Canvas(list_frame)
        scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=canvas.yview)
        subject_container = Frame(canvas)

        subject_container.bind(
            "<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=subject_container, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Populate subjects
        self._subject_container = subject_container
        for subject in subjects:
            self.create_subject_row(subject_container, subject)

        # Camera section: shows project default + optional override for this block.
        camera_frame = Frame(self)
        camera_frame.pack(fill="x", padx=20, pady=(0, 10))

        Label(
            camera_frame,
            text=_("📷 Camera:"),
            font=("Segoe UI", 10, "bold"),
        ).pack(side="left")

        self._camera_label = Label(
            camera_frame,
            text=self._format_current_camera(),
            font=("Segoe UI", 10),
            anchor="w",
        )
        self._camera_label.pack(side="left", padx=(5, 10))

        Button(
            camera_frame,
            text=_("Change..."),
            command=self._open_camera_chooser,
        ).pack(side="left")

        # Duration section: block-level default for this Day x Group.
        duration_frame = Frame(self)
        duration_frame.pack(fill="x", padx=20, pady=(0, 10))

        Label(
            duration_frame,
            text=_("⏱️ Default duration for the block:"),
            font=("Segoe UI", 10, "bold"),
        ).pack(side="left")

        self._duration_label = Label(
            duration_frame,
            text=self._format_block_duration(),
            font=("Segoe UI", 10),
            anchor="w",
        )
        self._duration_label.pack(side="left", padx=(5, 10))

        Button(
            duration_frame,
            text=_("Edit..."),
            command=self._edit_block_duration,
        ).pack(side="left")

        # Actions frame
        action_frame = Frame(self)
        action_frame.pack(fill="x", padx=20, pady=10)

        Label(
            action_frame,
            text=_("🛠️ Quick Actions"),
            font=("Segoe UI", 11, "bold"),
        ).pack(anchor="w", pady=(0, 10))

        Button(
            action_frame,
            text=_("▶️ Start Next Session"),
            command=self.start_next_session,
            width=30,
        ).pack(fill="x", pady=5)

        Button(
            action_frame,
            text=_("📊 Generate Partial Report"),
            command=self.generate_partial_report,
            width=30,
        ).pack(fill="x", pady=5)

        Button(
            action_frame,
            text=_("📝 Add Note"),
            command=self.add_note,
            width=30,
        ).pack(fill="x", pady=5)

        # Bottom buttons
        button_frame = Frame(self)
        button_frame.pack(fill="x", padx=20, pady=10)

        Button(
            button_frame,
            text=_("Close"),
            command=self.destroy,
        ).pack(side="right", padx=5)

        Button(
            button_frame,
            text=_("✅ Mark Batch as Complete"),
            command=self.mark_batch_complete,
        ).pack(side="right", padx=5)

    def _publish_project_views_refresh(self, reason: str) -> None:
        """Publish a project-views refresh if an event bus is available."""
        event_bus = getattr(self.session_coordinator, "event_bus", None)
        if event_bus is None:
            event_bus = getattr(self.live_batch_coordinator, "event_bus", None)
        if event_bus is None:
            return

        from zebtrack.ui import payloads
        from zebtrack.ui.event_bus_v2 import Event, UIEvents

        event_bus.publish(
            Event(
                type=UIEvents.PROJECT_VIEWS_REFRESH_REQUESTED,
                data=payloads.ProjectViewsRefreshRequestedPayload(
                    reason=reason,
                    append_summary=True,
                    immediate=True,
                ),
            )
        )

    @staticmethod
    def _open_generated_report_file(path: Path) -> None:
        """Open a generated report file using the platform opener."""
        from zebtrack.utils.os_opener import open_path

        open_path(path)

    def _find_session_folder(self, subject: str) -> Path | None:
        """Find the session folder for a specific day/group/subject.

        Lookup strategy (first match wins):
        1. ``project_data["batches"][*]["videos"][*].results_dir`` — preferred
           path that handles both legacy flat layouts and the new
           ``Grupo_X/Dia_Y/Sujeito_Z/live_{ts}/`` hierarchy uniformly.
        2. Legacy filesystem scan for ``day{N}_{group}_{subject}_*`` and
           ``D{N}_G{group}_S{subject}`` folders at the project root (used by
           older recordings made before the hierarchical layout was adopted).

        Args:
            subject: Subject ID (e.g., "1", "2")

        Returns:
            Path to session folder if found, None otherwise
        """
        if not self.project_manager.project_path:
            return None

        project_path = Path(self.project_manager.project_path)

        # Strategy 1 — read results_dir registered in project_data.
        results_dir = self._results_dir_for_subject(subject)
        if results_dir is not None:
            return results_dir

        # Strategy 2 — legacy filesystem scan at project root (pre-hierarchical
        # recordings whose results_dir was never stamped on the video entry).
        pattern_new = re.compile(
            rf"^day{self.day_num}_{re.escape(self.group_name)}_{subject}_\d{{8}}_\d{{6}}$"
        )
        pattern_legacy = re.compile(rf"^D{self.day_num}_G{re.escape(self.group_name)}_S{subject}$")

        for item in project_path.iterdir():
            if not item.is_dir():
                continue
            if pattern_new.match(item.name) or pattern_legacy.match(item.name):
                return item

        return None

    def _results_dir_for_subject(self, subject: str) -> Path | None:
        """Return the registered ``results_dir`` for the (day, group, subject) entry."""
        project_data = (
            self.project_manager.project_data
            if hasattr(self.project_manager, "project_data")
            else {}
        )
        target_day = f"Dia_{self.day_num}"
        for batch in project_data.get("batches", []):
            for video in batch.get("videos", []):
                metadata = video.get("metadata") or {}
                meta_day = str(metadata.get("day", "")).strip()
                if meta_day and meta_day not in (target_day, str(self.day_num)):
                    continue
                if str(metadata.get("group", "")).strip() != str(self.group_name).strip():
                    continue
                if str(metadata.get("subject", "")).strip() != str(subject).strip():
                    continue
                results_dir = video.get("results_dir")
                if results_dir:
                    candidate = Path(results_dir)
                    if candidate.exists() and candidate.is_dir():
                        return candidate
        return None

    def _get_polygon_source_for_subject(self, subject: str) -> str | None:
        """Return the polygon source ("auto" / "manual" / None) recorded for a subject.

        Scans the project's videos for an entry whose metadata matches the
        block's (group, day, subject) tuple. Returns the ``polygon_source``
        field stamped by ``OutputRegistrationManager.register_processing_outputs``
        after the live recording completes, or ``None`` for sessions recorded
        before the field existed.
        """
        project_data = (
            self.project_manager.project_data
            if hasattr(self.project_manager, "project_data")
            else {}
        )
        target_day = f"Dia_{self.day_num}"
        for batch in project_data.get("batches", []):
            for video in batch.get("videos", []):
                metadata = video.get("metadata") or {}
                # Day field uses both "Dia_N" and bare int formats across the codebase
                meta_day = str(metadata.get("day", "")).strip()
                if meta_day and meta_day not in (target_day, str(self.day_num)):
                    continue
                if str(metadata.get("group", "")).strip() != str(self.group_name).strip():
                    continue
                if str(metadata.get("subject", "")).strip() != str(subject).strip():
                    continue
                source = metadata.get("polygon_source")
                if source:
                    return str(source)
        return None

    def _get_session_files_status(self, folder: Path) -> dict[str, bool]:
        """Check which output files exist in a session folder.

        Args:
            folder: Path to session folder

        Returns:
            Dict with file type as key and existence as value
        """
        status = {
            "video": False,
            "trajectory": False,  # 3_CoordMovimento
            "arena": False,  # 1_ProcessingArea
            "rois": False,  # 2_AreasOfInterest or 2_ZonasROI
            "summary": False,  # 4_Resumo or similar
        }

        if not folder or not folder.exists():
            return status

        for file in folder.iterdir():
            name = file.name.lower()
            if file.suffix == ".mp4":
                status["video"] = True
            elif "coordmovimento" in name or "trajectory" in name:
                status["trajectory"] = True
            elif "processingarea" in name or "arena" in name:
                status["arena"] = True
            elif "areasofinterest" in name or "zonasroi" in name or "zonas" in name:
                status["rois"] = True

        status["summary"] = has_summary_excel_output(folder)

        return status

    def create_subject_row(self, parent: Frame, subject: str):
        """Create row for single subject.

        Args:
            parent: Parent frame to contain the row
            subject: Subject ID (e.g., "1", "2", etc.)
        """
        # v2.3.1: Use day_num (int) for session lookup
        is_completed = (self.day_num, self.group_name, subject) in self.completed_sessions

        # v2.3.1: Get session folder and file status
        session_folder = self._find_session_folder(subject)
        files_status = self._get_session_files_status(session_folder) if session_folder else {}

        row = Frame(parent, relief="solid", borderwidth=1, bg="white")
        row.pack(fill="x", padx=5, pady=3)

        # Status indicator
        if is_completed:
            status_label = Label(row, text="✅", font=("Segoe UI", 14), bg="white")
            status_text = _("Recorded")
        else:
            status_label = Label(row, text="⏸️", font=("Segoe UI", 14), bg="white")
            # NOT _("Pending"): that msgid is the status-legend label shared by
            # project_overview/processing_reports/validation_manager, whose
            # pt_BR text is the plural "Pendentes" (it counts items). This is a
            # single subject's row, so it needs its own msgid.
            status_text = _("Not recorded")

        status_label.pack(side="left", padx=10, pady=10)

        # Subject info
        info_frame = Frame(row, bg="white")
        info_frame.pack(side="left", fill="both", expand=True, padx=10, pady=10)

        Label(
            info_frame,
            text=_("Animal {subject}").format(subject=subject),
            font=("Segoe UI", 11, "bold"),
            bg="white",
        ).pack(anchor="w")

        # v2.3.1: Show file status icons for completed sessions
        if is_completed and files_status:
            file_icons = []
            if files_status.get("video"):
                file_icons.append("🎬")  # Video
            if files_status.get("arena"):
                file_icons.append("🏟️")  # Arena
            if files_status.get("trajectory"):
                file_icons.append("🧭")  # Trajectory
            if files_status.get("rois"):
                file_icons.append("🎯")  # ROIs
            if files_status.get("summary"):
                file_icons.append("Σ")  # Summary

            files_text = " ".join(file_icons) if file_icons else _("⚠️ No files")
            status_detail = f"{status_text} | {files_text}"
        else:
            status_detail = status_text

        Label(
            info_frame,
            text=status_detail,
            font=("Segoe UI", 9),
            fg="#666",
            bg="white",
        ).pack(anchor="w")

        # Polygon-source badge: shows whether the polygon used for this
        # subject was auto-detected or manually drawn. For completed sessions
        # we read the value stamped by ``register_processing_outputs``; for
        # pending subjects we hint that the project polygon will be reused.
        if is_completed:
            polygon_source = self._get_polygon_source_for_subject(subject)
            if polygon_source == "auto":
                Label(
                    info_frame,
                    text=_("🏟️ Auto-detected"),
                    font=("Segoe UI", 8, "bold"),
                    fg="#0a7",
                    bg="white",
                ).pack(anchor="w")
            elif polygon_source == "manual":
                Label(
                    info_frame,
                    text=_("✏️ Drawn manually"),
                    font=("Segoe UI", 8, "bold"),
                    fg="#666",
                    bg="white",
                ).pack(anchor="w")
        elif self._project_has_polygon:
            Label(
                info_frame,
                text=_("🏟️ Project polygon ready (will be reused)"),
                font=("Segoe UI", 8),
                fg="#0a7",
                bg="white",
            ).pack(anchor="w")

        # Duração desta cobaia. Para pendentes é o que SERÁ usado (e é editável);
        # para gravadas é registro do que foi planejado — a duração real do vídeo
        # vive no resumo da sessão, e reescrever o override depois do fato só
        # falsificaria o histórico.
        project_data = self._project_data()
        duration_s = resolve_session_duration(project_data, self.day_num, self.group_name, subject)
        overrides = project_data.get("session_duration_overrides") or {}
        from zebtrack.core.services.session_duration_resolver import duration_override_key

        has_own = duration_override_key(self.day_num, self.group_name, subject) in overrides
        Label(
            info_frame,
            text=(
                _("⏱️ {duration} (its own)").format(duration=self._format_duration(duration_s))
                if has_own
                else _("⏱️ {duration} (block default)").format(
                    duration=self._format_duration(duration_s)
                )
            ),
            font=("Segoe UI", 8, "bold" if has_own else "normal"),
            fg="#b36b00" if has_own else "#666",
            bg="white",
        ).pack(anchor="w")

        # v2.3.1: Show folder name if exists
        if session_folder:
            Label(
                info_frame,
                text=f"📁 {session_folder.name}",
                font=("Segoe UI", 8),
                fg="#999",
                bg="white",
            ).pack(anchor="w")

        # Action buttons
        if is_completed:
            Button(
                row,
                text=_("📊 View Results"),
                command=lambda: self.view_results(subject),
            ).pack(side="right", padx=5, pady=10)
        else:
            Button(
                row,
                text=_("▶️ Start"),
                command=lambda: self.start_session(subject),
            ).pack(side="right", padx=5, pady=10)
            Button(
                row,
                text=_("⏱️ Duration"),
                command=lambda: self._edit_subject_duration(subject),
            ).pack(side="right", padx=5, pady=10)

    def _format_current_camera(self) -> str:
        """Render the camera label: override (if set) or project default."""
        if self._camera_index_override is not None:
            name = self._camera_friendly_name_override or ""
            suffix = f" — {name}" if name else ""
            return _("[Session] Index {index}{suffix}").format(
                index=self._camera_index_override, suffix=suffix
            )

        project_data = (
            self.project_manager.project_data
            if hasattr(self.project_manager, "project_data")
            else {}
        )
        saved_index = project_data.get("camera_index", 0)
        saved_name = project_data.get("camera_friendly_name", "") or ""
        if saved_name:
            return _("{name} (index {index})").format(name=saved_name, index=saved_index)
        return _("Index {index}").format(index=saved_index)

    def _open_camera_chooser(self) -> None:
        """Modal sub-dialog: detect + pick a camera (and optionally persist)."""
        # Local imports keep dialog cold-import light.
        from tkinter import BooleanVar, Checkbutton, StringVar

        from zebtrack.core.services.wizard_service import WizardService

        try:
            cameras = WizardService.detect_available_cameras(use_cache=False)
        # except Exception justified: camera enumeration is hardware I/O
        except Exception as exc:
            messagebox.showerror(
                _("Detection failed"),
                _("Could not detect cameras:\n\n{error}").format(error=exc),
                parent=self,
            )
            return

        if not cameras:
            messagebox.showwarning(
                _("No camera"),
                _("No camera was detected on this system."),
                parent=self,
            )
            return

        chooser = Toplevel(self)
        chooser.title(_("Change the camera for this session"))
        chooser.transient(self)
        chooser.grab_set()

        Label(chooser, text=_("Select the camera:")).grid(
            row=0, column=0, sticky="w", padx=10, pady=(10, 0)
        )

        descriptions = [
            c.get("description", _("Camera {index}").format(index=c["index"])) for c in cameras
        ]
        index_map = {desc: int(cameras[i]["index"]) for i, desc in enumerate(descriptions)}
        name_map = {
            desc: cameras[i].get("friendly_name", "") for i, desc in enumerate(descriptions)
        }

        selection_var = StringVar(value=descriptions[0])
        combo = ttk.Combobox(
            chooser,
            values=descriptions,
            textvariable=selection_var,
            state="readonly",
            width=60,
        )
        combo.grid(row=1, column=0, columnspan=2, sticky="ew", padx=10, pady=5)

        persist_var = BooleanVar(value=False)
        Checkbutton(
            chooser,
            text=_("Save as the default camera for this project"),
            variable=persist_var,
        ).grid(row=2, column=0, columnspan=2, sticky="w", padx=10, pady=5)

        confirmed = {"ok": False}

        def _on_ok() -> None:
            confirmed["ok"] = True
            chooser.destroy()

        def _on_cancel() -> None:
            chooser.destroy()

        button_row = Frame(chooser)
        button_row.grid(row=3, column=0, columnspan=2, pady=(5, 10))
        Button(button_row, text="OK", command=_on_ok, width=10).pack(side="left", padx=5)
        Button(button_row, text=_("Cancel"), command=_on_cancel, width=10).pack(side="left", padx=5)

        chooser.wait_window()

        if not confirmed["ok"]:
            return

        chosen = selection_var.get()
        new_index = index_map.get(chosen, 0)
        new_name = name_map.get(chosen, "")

        if persist_var.get():
            try:
                self.project_manager.project_data["camera_index"] = int(new_index)
                self.project_manager.project_data["camera_friendly_name"] = new_name
                if hasattr(self.project_manager, "save_project"):
                    self.project_manager.save_project()
                # Persisted: clear any previous override so the label shows the new default.
                self._camera_index_override = None
                self._camera_friendly_name_override = None
            except (OSError, AttributeError, ValueError) as exc:
                messagebox.showwarning(
                    _("Failed to save camera"),
                    _("Could not save the camera as the default:\n{error}").format(error=exc),
                    parent=self,
                )
                # Fall back to per-session override on save failure.
                self._camera_index_override = int(new_index)
                self._camera_friendly_name_override = new_name
        else:
            self._camera_index_override = int(new_index)
            self._camera_friendly_name_override = new_name

        if self._camera_label is not None:
            self._camera_label.config(text=self._format_current_camera())

    # ------------------------------------------------------------------
    # Recording duration (block default + per-subject override)
    # ------------------------------------------------------------------

    def _project_data(self) -> dict:
        """Return ``project_data``, or an empty dict when unavailable."""
        return getattr(self.project_manager, "project_data", None) or {}

    @staticmethod
    def _format_duration(duration_s: float) -> str:
        """Render seconds as the minutes the operator actually thinks in."""
        minutes = duration_s / 60.0
        if abs(minutes - round(minutes)) < 0.01:
            return f"{round(minutes)} min"
        return f"{minutes:.1f} min"

    def _format_block_duration(self) -> str:
        """Label for the block default, flagging where the value came from."""
        project_data = self._project_data()
        # O bloco não tem "cobaia", então resolvemos com o curinga: isso pula o
        # nível de cobaia e cai direto em bloco > projeto.
        duration = resolve_session_duration(
            project_data, self.day_num, self.group_name, SUBJECT_WILDCARD
        )
        overrides = project_data.get("session_duration_overrides") or {}
        from zebtrack.core.services.session_duration_resolver import block_override_key

        has_block_override = block_override_key(self.day_num, self.group_name) in overrides
        origin = "bloco" if has_block_override else "projeto"
        return f"{self._format_duration(duration)} ({origin})"

    def _ask_duration_minutes(self, title: str, prompt: str, current_s: float) -> float | None:
        """Prompt for a duration in minutes. Returns seconds, or None if cancelled."""
        answer = simpledialog.askstring(
            title,
            prompt,
            initialvalue=f"{current_s / 60.0:g}",
            parent=self,
        )
        if answer is None:
            return None

        try:
            minutes = float(answer.strip().replace(",", "."))
        except ValueError:
            messagebox.showwarning(
                _("Invalid value"),
                _("'{value}' is not a number of minutes.").format(value=answer),
                parent=self,
            )
            return None

        if minutes <= 0:
            messagebox.showwarning(
                _("Invalid value"),
                _("The duration must be greater than zero."),
                parent=self,
            )
            return None

        return minutes * 60.0

    def _persist_project(self) -> bool:
        """Save the project, surfacing failures instead of swallowing them.

        ``save_project()`` raises when ``project_path`` is unset, and callers
        reached through the event bus have that exception eaten by the bus's
        try/except — the user would see the new duration on screen and lose it
        on restart. Check and report here.
        """
        if not getattr(self.project_manager, "project_path", None):
            messagebox.showerror(
                _("Project not saved"),
                _(
                    "The project has no path defined, so the duration could not "
                    "be written. Save the project and try again."
                ),
                parent=self,
            )
            return False

        try:
            if hasattr(self.project_manager, "save_project"):
                self.project_manager.save_project()
        # except Exception justified: qualquer falha de I/O aqui precisa virar
        # feedback honesto — silenciar faria a duração "sumir" no próximo boot.
        except Exception as exc:
            log.error("block_detail.duration.save_failed", error=str(exc), exc_info=True)
            messagebox.showerror(
                _("Failed to save"),
                _("The duration could not be written to the project:\n{error}").format(error=exc),
                parent=self,
            )
            return False

        return True

    def _edit_block_duration(self) -> None:
        """Set/clear the Day x Group default duration."""
        project_data = self._project_data()
        current = resolve_session_duration(
            project_data, self.day_num, self.group_name, SUBJECT_WILDCARD
        )

        new_duration = self._ask_duration_minutes(
            _("Default duration for the block"),
            _(
                "Recording duration for {block}, in minutes:\n\n"
                "Applies to every subject in this block that has no duration of "
                "its own. Sessions already recorded are not affected."
            ).format(block=_block_label(self.day_num, self.group_name)),
            current,
        )
        if new_duration is None:
            return

        set_duration_override(
            project_data, self.day_num, self.group_name, SUBJECT_WILDCARD, new_duration
        )
        if not self._persist_project():
            return

        if self._duration_label is not None:
            self._duration_label.config(text=self._format_block_duration())
        self.refresh_subject_rows()

    def _edit_subject_duration(self, subject: str) -> None:
        """Set/clear the per-subject duration override."""
        project_data = self._project_data()
        current = resolve_session_duration(project_data, self.day_num, self.group_name, subject)
        block_default = resolve_session_duration(
            project_data, self.day_num, self.group_name, SUBJECT_WILDCARD
        )

        new_duration = self._ask_duration_minutes(
            _("Duration — Animal {subject}").format(subject=subject),
            _(
                "Recording duration for Animal {subject} "
                "({block}), in minutes:\n\n"
                "Leave it equal to {default} to follow the block default."
            ).format(
                subject=subject,
                block=_block_label(self.day_num, self.group_name),
                default=self._format_duration(block_default),
            ),
            current,
        )
        if new_duration is None:
            return

        # Igualar ao padrão do bloco = voltar a herdar. Guardar um override
        # idêntico ao pai só criaria ruído no JSON e mentiria na UI ("próprio"
        # quando na verdade é herdado).
        if abs(new_duration - block_default) < 0.5:
            set_duration_override(project_data, self.day_num, self.group_name, subject, None)
        else:
            set_duration_override(
                project_data, self.day_num, self.group_name, subject, new_duration
            )

        if not self._persist_project():
            return

        self.refresh_subject_rows()

    def refresh_subject_rows(self) -> None:
        """Rebuild the subject list so duration labels reflect the new values."""
        container = getattr(self, "_subject_container", None)
        if container is None:
            return
        for child in container.winfo_children():
            child.destroy()
        for subject in self._subjects():
            self.create_subject_row(container, subject)

    def _subjects(self) -> list[str]:
        """Subject IDs of this block, as strings ("1", "2", ...)."""
        return [str(i + 1) for i in range(self.subjects_per_group)]

    def start_session(self, subject: str):
        """Start live session for subject.

        Args:
            subject: Subject ID to start session for
        """
        # Resolver ANTES do destroy: depois de fechar o diálogo ainda dá para ler
        # project_manager, mas manter a leitura aqui deixa o valor no log junto
        # do resto do contexto da sessão.
        duration_s = resolve_session_duration(
            self._project_data(), self.day_num, self.group_name, subject
        )

        log.info(
            "block_detail.start_session",
            day=self.day_num,
            group=self.group_name,
            subject=subject,
            camera_override=self._camera_index_override,
            duration_s=duration_s,
        )

        # v2.3.1: Actually start the session using session_coordinator
        try:
            # Snapshot the override before destroying (instance attrs survive, but be explicit).
            override_index = self._camera_index_override
            override_name = self._camera_friendly_name_override

            # Close dialog first so it doesn't block
            self.destroy()

            # Start the live project session
            success = self.session_coordinator.start_live_project_session(
                day=self.day_num,
                group=str(self.group_name),
                subject=subject,
                duration_s=duration_s,
                camera_index_override=override_index,
                camera_friendly_name_override=override_name,
            )

            if not success:
                # ``start_live_project_session`` returns False in three distinct
                # situations: a genuine failure (camera missing, bad project
                # type) AND two legitimate deferrals — "awaiting zone
                # confirmation" after the auto-detect flow approves a polygon,
                # and "armed, awaiting the external Arduino trigger". In both
                # deferred cases the user already has the right affordance on
                # screen (the LIVE_RECORDING_PENDING banner, or the "Aguardando
                # sinal externo" notice) — an error popup here would be a lie.
                cal_coord = getattr(self.session_coordinator, "live_calibration_coordinator", None)
                deferred = bool(
                    cal_coord is not None and getattr(cal_coord, "pending_zone_confirmation", False)
                )
                if not deferred and hasattr(
                    self.session_coordinator, "has_pending_external_trigger"
                ):
                    deferred = bool(self.session_coordinator.has_pending_external_trigger())
                if deferred:
                    log.info(
                        "block_detail.start_session.deferred_for_zones",
                        subject=subject,
                        day=self.day_num,
                        group=self.group_name,
                    )
                else:
                    messagebox.showerror(
                        _("Error"),
                        _("Failed to start the session for Animal {subject}\n{block}").format(
                            subject=subject,
                            block=_block_label(self.day_num, self.group_name),
                        ),
                    )
        except Exception as e:
            log.error("block_detail.start_session.failed", error=str(e), exc_info=True)
            messagebox.showerror(
                _("Error"), _("Error starting the session: {error}").format(error=e)
            )

    def start_next_session(self):
        """Start next pending session."""
        # v2.3.1: Use subjects_per_group and completed_sessions
        subjects = self._subjects()
        for subject in subjects:
            if (self.day_num, self.group_name, subject) not in self.completed_sessions:
                self.start_session(subject)
                return

        messagebox.showinfo(_("Complete"), _("Every session in this block has been completed!"))

    def view_results(self, subject: str):
        """View session results by opening the session folder.

        Args:
            subject: Subject ID to view results for
        """
        session_folder = self._find_session_folder(subject)

        if not session_folder or not session_folder.exists():
            messagebox.showwarning(
                _("Folder not found"),
                _("Could not find the results folder for Animal {subject}.\n{block}").format(
                    subject=subject, block=_block_label(self.day_num, self.group_name)
                ),
            )
            return

        try:
            # Open folder in system file explorer
            from zebtrack.utils.os_opener import open_path

            open_path(session_folder)

            log.info(
                "block_detail.view_results.opened",
                folder=str(session_folder),
            )
        except Exception as e:
            log.error("block_detail.view_results.failed", error=str(e), exc_info=True)
            messagebox.showerror(
                _("Error"),
                _("Failed to open the results folder:\n{error}").format(error=str(e)),
            )

    def _get_completed_subjects_for_partial_report(self) -> list[str]:
        subjects = self._subjects()
        return [
            subject
            for subject in subjects
            if (self.day_num, self.group_name, subject) in self.completed_sessions
        ]

    def _collect_partial_report_summary_files(
        self, completed_subjects: list[str]
    ) -> list[tuple[str, Path]]:
        summary_files = []
        for subject in completed_subjects:
            session_folder = self._find_session_folder(subject)
            if not session_folder or not session_folder.exists():
                continue

            summary_file = find_summary_excel_file(session_folder)
            if summary_file is not None:
                summary_files.append((subject, summary_file))

        return summary_files

    def _build_partial_report_dataset(self, summary_files: list[tuple[str, Path]]):
        import warnings

        import pandas as pd

        all_data = []
        parsed_summary_files = []
        for subject, summary_path in summary_files:
            try:
                df = pd.read_excel(summary_path)
                df["animal"] = subject
                df["dia"] = self.day_num
                df["grupo"] = self.group_name
                df["source_file"] = summary_path.name
                all_data.append(df)
                parsed_summary_files.append((subject, summary_path))
            except Exception as e:
                log.warning(
                    "block_detail.partial_report.read_failed",
                    summary_path=str(summary_path),
                    error=str(e),
                )

        if not all_data:
            raise ValueError(_("No valid data found in the summary files"))

        non_empty_dfs = [df for df in all_data if not df.empty]
        if not non_empty_dfs:
            raise ValueError(_("No valid data found (every file was empty)"))

        with warnings.catch_warnings():
            warnings.filterwarnings(
                "ignore",
                category=FutureWarning,
                message=".*concatenation with empty or all-NA entries.*",
            )
            unified_df = pd.concat(non_empty_dfs, ignore_index=True)

        return all_data, unified_df, parsed_summary_files

    @staticmethod
    def _get_partial_report_stats_columns(unified_df) -> list[str]:
        import pandas as pd

        # ``duration`` entra explicitamente: nenhum dos outros keywords o captura
        # (a coluna é ``video_duration_s``), e sem ela o agregado esconde
        # justamente o dado que torna as métricas absolutas comparáveis — ou não.
        keywords = ["distance", "speed", "time", "entries", "duration"]
        return [
            col
            for col in unified_df.columns
            if col != "analysis_timestamp"
            and any(keyword in col.lower() for keyword in keywords)
            and pd.api.types.is_numeric_dtype(unified_df[col])
        ]

    def _heterogeneous_duration_warning(self, subjects: list[str]) -> str | None:
        """Aviso quando as cobaias do bloco não compartilham a mesma duração.

        Métricas ABSOLUTAS (distância total, nº de entradas, tempo em ROI) crescem
        com o tempo de gravação. Agregá-las por média entre animais gravados por
        tempos diferentes produz um número que parece comparável e não é. O app
        não normaliza sozinho — essa decisão é do pesquisador — mas também não
        cala sobre o problema.
        """
        durations = collect_block_durations(
            self._project_data(), self.day_num, self.group_name, subjects
        )
        distinct = sorted({round(value, 3) for value in durations.values()})
        if len(distinct) <= 1:
            return None

        listed = ", ".join(self._format_duration(value) for value in distinct)
        return _(
            "The sessions in this block have different durations ({listed}).\n\n"
            "ABSOLUTE metrics — total distance, number of entries, time in ROI — "
            "grow with recording time and are NOT directly comparable across "
            "these animals. The 'video_duration_s' column is in the report so you "
            "can normalize however you prefer.\n\n"
            "Generate the report anyway?"
        ).format(listed=listed)

    @staticmethod
    def _format_partial_report_cell_value(value) -> str:
        import pandas as pd

        if pd.isna(value):
            return "-"
        if isinstance(value, int):
            return str(value)
        if isinstance(value, float):
            return str(int(value)) if value.is_integer() else f"{value:.2f}"
        return str(value)

    @staticmethod
    def _build_partial_report_output_paths(
        reports_dir: Path, base_output_name: str
    ) -> tuple[str, Path, str, Path]:
        excel_output_name = f"{base_output_name}.xlsx"
        word_output_name = f"{base_output_name}.docx"
        return (
            excel_output_name,
            reports_dir / excel_output_name,
            word_output_name,
            reports_dir / word_output_name,
        )

    def _write_partial_report_excel(self, path: Path, all_data, unified_df) -> None:
        import pandas as pd

        with pd.ExcelWriter(path, engine="openpyxl") as writer:
            # Sheet names are a persistence contract, like the por_animal sheet:
            # third-party scripts open these workbooks by sheet name, so they
            # must not change with ui.language.
            unified_df.to_excel(writer, sheet_name="Dados Consolidados", index=False)

            stats_cols = self._get_partial_report_stats_columns(unified_df)
            if len(all_data) > 1 and stats_cols:
                summary_stats = unified_df.groupby("animal")[stats_cols].mean()
                summary_stats.to_excel(writer, sheet_name="Resumo por Animal")  # i18n: not-ui

    def _write_partial_report_word(
        self,
        path: Path,
        excel_name: str,
        parsed_summary_files: list[tuple[str, Path]],
        all_data,
        unified_df,
    ) -> None:
        from docx import Document

        document = Document()
        document.add_heading(
            _("Partial Report - {block}").format(block=_block_label(self.day_num, self.group_name)),
            level=1,
        )
        document.add_paragraph(
            _("Generated at: {timestamp}").format(
                timestamp=datetime.now().strftime("%d/%m/%Y %H:%M:%S")
            )
        )
        document.add_paragraph(_("Aggregated sessions: {count}").format(count=len(all_data)))
        document.add_paragraph(_("Consolidated spreadsheet: {name}").format(name=excel_name))

        # A ressalva vai para DENTRO do documento, não só para o popup: quem lê o
        # relatório meses depois não viu a caixa de diálogo.
        subjects_in_report = [subject for subject, _path in parsed_summary_files]
        durations = collect_block_durations(
            self._project_data(), self.day_num, self.group_name, subjects_in_report
        )
        distinct = sorted({round(value, 3) for value in durations.values()})
        if len(distinct) > 1:
            listed = ", ".join(self._format_duration(value) for value in distinct)
            document.add_paragraph(
                _(
                    "WARNING — heterogeneous recording durations in this block "
                    "({listed}). Absolute metrics (total distance, number of "
                    "entries, time in ROI) scale with recording time and are not "
                    "directly comparable across these animals without "
                    "normalization. Use the 'video_duration_s' column of the "
                    "spreadsheet to normalize."
                ).format(listed=listed)
            )

        document.add_heading(_("Sessions included"), level=2)
        session_table = document.add_table(rows=1, cols=2)
        session_table.style = "Table Grid"
        session_table.rows[0].cells[0].text = _("Animal")
        session_table.rows[0].cells[1].text = _("Source file")

        for subject, summary_path in parsed_summary_files:
            row_cells = session_table.add_row().cells
            row_cells[0].text = str(subject)
            row_cells[1].text = summary_path.name

        stats_cols = self._get_partial_report_stats_columns(unified_df)
        if stats_cols:
            summary_stats = unified_df.groupby("animal")[stats_cols].mean().reset_index()
            document.add_heading(_("Summary per Animal"), level=2)
            summary_table = document.add_table(rows=1, cols=len(summary_stats.columns))
            summary_table.style = "Table Grid"
            header_cells = summary_table.rows[0].cells
            for idx, column_name in enumerate(summary_stats.columns):
                header_cells[idx].text = str(column_name)

            for _idx, row_data in summary_stats.iterrows():
                row_cells = summary_table.add_row().cells
                for idx, column_name in enumerate(summary_stats.columns):
                    row_cells[idx].text = self._format_partial_report_cell_value(
                        row_data[column_name]
                    )

        document.save(str(path))

    def _write_partial_report_outputs(
        self,
        reports_dir: Path,
        parsed_summary_files: list[tuple[str, Path]],
        all_data,
        unified_df,
    ) -> tuple[str, Path, str, Path, bool]:
        base_output_name = f"PartialReport_Dia{self.day_num}_{self.group_name}"
        (
            excel_output_name,
            excel_output_path,
            word_output_name,
            word_output_path,
        ) = self._build_partial_report_output_paths(reports_dir, base_output_name)

        write_fallback_used = False
        try:
            self._write_partial_report_excel(excel_output_path, all_data, unified_df)
            self._write_partial_report_word(
                word_output_path,
                excel_output_name,
                parsed_summary_files,
                all_data,
                unified_df,
            )
        except PermissionError:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            fallback_output_name = f"{base_output_name}_{timestamp}"
            (
                excel_output_name,
                excel_output_path,
                word_output_name,
                word_output_path,
            ) = self._build_partial_report_output_paths(reports_dir, fallback_output_name)
            self._write_partial_report_excel(excel_output_path, all_data, unified_df)
            self._write_partial_report_word(
                word_output_path,
                excel_output_name,
                parsed_summary_files,
                all_data,
                unified_df,
            )
            write_fallback_used = True

        return (
            excel_output_name,
            excel_output_path,
            word_output_name,
            word_output_path,
            write_fallback_used,
        )

    def _notify_partial_report_success(
        self,
        excel_output_name: str,
        word_output_name: str,
        session_count: int,
        write_fallback_used: bool,
    ) -> None:
        if write_fallback_used:
            messagebox.showwarning(
                _("File in use"),
                _(
                    "The default file was locked by another program/sync "
                    "service.\n"
                    "The reports were saved under new names:\n"
                    "{excel}\n{word}"
                ).format(excel=excel_output_name, word=word_output_name),
            )

        messagebox.showinfo(
            _("Reports Generated"),
            _(
                "Partial reports generated successfully!\n\n"
                "📊 Excel: {excel}\n"
                "📝 Word: {word}\n"
                "🐟 {count} aggregated sessions"
            ).format(excel=excel_output_name, word=word_output_name, count=session_count),
        )

    def _prompt_open_partial_report_files(
        self,
        excel_output_path: Path,
        excel_output_name: str,
        word_output_path: Path,
        word_output_name: str,
    ) -> None:
        if messagebox.askyesno(
            _("Open Partial Report"),
            _("Open the partial spreadsheet in Excel?\n\n📊 {name}").format(name=excel_output_name),
        ):
            try:
                self._open_generated_report_file(excel_output_path)
            except Exception as e:
                log.warning("block_detail.partial_report.open_failed", error=str(e))
                messagebox.showwarning(
                    _("Warning"),
                    _("The Excel report was generated, but it could not be opened:\n{path}").format(
                        path=excel_output_path
                    ),
                )

        if messagebox.askyesno(
            _("Open Partial Report"),
            _("Open the partial report in Word?\n\n📝 {name}").format(name=word_output_name),
        ):
            try:
                self._open_generated_report_file(word_output_path)
            except Exception as e:
                log.warning("block_detail.partial_report.open_failed", error=str(e))
                messagebox.showwarning(
                    _("Warning"),
                    _("The Word report was generated, but it could not be opened:\n{path}").format(
                        path=word_output_path
                    ),
                )

    def generate_partial_report(self):
        """Generate partial report for completed sessions in this block.

        Collects all summary Excel files from completed sessions and aggregates
        them into Excel and Word outputs for the day/group block.
        """
        log.info(
            "block_detail.generate_partial_report",
            day=self.day_num,
            group=self.group_name,
        )

        completed_in_block = self._get_completed_subjects_for_partial_report()

        if not completed_in_block:
            messagebox.showwarning(
                _("No Sessions"),
                _("No completed session found for\n{block}").format(
                    block=_block_label(self.day_num, self.group_name)
                ),
            )
            return

        summary_files = self._collect_partial_report_summary_files(completed_in_block)

        if not summary_files:
            messagebox.showwarning(
                _("No Reports"),
                _(
                    "No summary file found in the sessions of\n{block}\n\n"
                    "Run the session analysis first."
                ).format(block=_block_label(self.day_num, self.group_name)),
            )
            return

        warning = self._heterogeneous_duration_warning(completed_in_block)
        if warning and not messagebox.askyesno(_("Different durations in the block"), warning):
            log.info(
                "block_detail.partial_report.cancelled_on_duration_warning",
                day=self.day_num,
                group=self.group_name,
            )
            return

        try:
            reports_dir = Path(self.project_manager.project_path) / "partial_reports"
            reports_dir.mkdir(parents=True, exist_ok=True)

            all_data, unified_df, parsed_summary_files = self._build_partial_report_dataset(
                summary_files
            )
            (
                excel_output_name,
                excel_output_path,
                word_output_name,
                word_output_path,
                write_fallback_used,
            ) = self._write_partial_report_outputs(
                reports_dir,
                parsed_summary_files,
                all_data,
                unified_df,
            )

            log.info(
                "block_detail.partial_report.success",
                excel_output=str(excel_output_path),
                word_output=str(word_output_path),
                session_count=len(all_data),
            )

            self._publish_project_views_refresh(
                _("Partial reports updated: {block}").format(
                    block=_block_label(self.day_num, self.group_name)
                )
            )

            self._notify_partial_report_success(
                excel_output_name,
                word_output_name,
                len(all_data),
                write_fallback_used,
            )
            self._prompt_open_partial_report_files(
                excel_output_path,
                excel_output_name,
                word_output_path,
                word_output_name,
            )

        except Exception as e:
            log.error("block_detail.generate_partial_report.failed", error=str(e), exc_info=True)
            messagebox.showerror(
                _("Error"),
                _("Failed to generate the partial report:\n{error}").format(error=e),
            )

    def add_note(self):
        """Add note to day/group block."""
        log.info("block_detail.add_note", day=self.day_num, group=self.group_name)

        # Get existing notes from project data
        project_data = (
            self.project_manager.project_data
            if hasattr(self.project_manager, "project_data")
            else {}
        )

        # Notes are stored in experiment_notes dict with block key
        experiment_notes = project_data.get("experiment_notes", {})
        block_key = f"Dia_{self.day_num}_{self.group_name}"
        existing_note = experiment_notes.get(block_key, "")

        # Show input dialog
        note = simpledialog.askstring(
            _("Add Experimental Note"),
            _("Note for {block}:\n\n(leave blank to clear)").format(
                block=_block_label(self.day_num, self.group_name)
            ),
            initialvalue=existing_note,
            parent=self,
        )

        if note is None:
            # User cancelled
            return

        try:
            # Save note
            if "experiment_notes" not in project_data:
                project_data["experiment_notes"] = {}

            if note.strip():
                project_data["experiment_notes"][block_key] = note.strip()
                log.info("block_detail.add_note.saved", block_key=block_key, note=note[:50])
                messagebox.showinfo(_("Note Saved"), _("Experimental note saved successfully!"))
            else:
                # Remove note if empty
                if block_key in project_data["experiment_notes"]:
                    del project_data["experiment_notes"][block_key]
                log.info("block_detail.add_note.cleared", block_key=block_key)
                messagebox.showinfo(_("Note Removed"), _("Experimental note removed."))

            # Save project
            if hasattr(self.project_manager, "save_project"):
                self.project_manager.save_project()

        except Exception as e:
            log.error("block_detail.add_note.failed", error=str(e), exc_info=True)
            messagebox.showerror(_("Error"), _("Failed to save the note:\n{error}").format(error=e))

    def mark_batch_complete(self):
        """Mark batch as complete and generate the block partial report.

        Reusa o mesmo gerador do botão "Gerar Relatório Parcial" (Excel +
        Word em ``partial_reports/``), rodando em thread de fundo para
        honrar a promessa de "segundo plano", e persiste a completude do
        lote via ``LiveBatchCoordinator.mark_block_complete`` — o que pinta
        o quadrado do grid de verde e sobrevive ao reinício do app. O
        caminho antigo montava um ``batch_id`` com ``*`` literal que nunca
        casava com os IDs reais e ignorava o retorno, mostrando sucesso sem
        gerar nada.
        """
        # Audit Erro 6 (2026-05-25): make the scope explicit so the user
        # knows exactly which sessions are being consolidated. A "lote" here
        # is all sessions of THIS group on THIS day (one row of the grid),
        # not the whole project.
        result = messagebox.askyesno(
            _("Confirm — Mark batch as complete"),
            _(
                "Mark the batch of Group '{group}' on Day {day} as complete?\n\n"
                "Scope: ALL sessions already recorded for this group on this day "
                "will be consolidated into the block's partial report (Excel + "
                "Word) and the matching square in the Progress grid will turn "
                "green.\n\n"
                "This action does NOT affect other groups, other days, nor does it "
                "close the project as a whole. You can carry on recording new "
                "subjects on other days/groups normally.\n\n"
                "Continue?"
            ).format(group=self.group_name, day=self.day_num),
        )
        if not result:
            return

        # Pré-checagens rápidas na thread da UI — feedback honesto quando não
        # há sessões ou resumos (mesmos textos do "Gerar Relatório Parcial").
        completed_in_block = self._get_completed_subjects_for_partial_report()
        if not completed_in_block:
            messagebox.showwarning(
                _("No Sessions"),
                _("No completed session found for\n{block}").format(
                    block=_block_label(self.day_num, self.group_name)
                ),
            )
            return

        summary_files = self._collect_partial_report_summary_files(completed_in_block)
        if not summary_files:
            messagebox.showwarning(
                _("No Reports"),
                _(
                    "No summary file found in the sessions of\n{block}\n\n"
                    "Run the session analysis first."
                ).format(block=_block_label(self.day_num, self.group_name)),
            )
            return

        warning = self._heterogeneous_duration_warning(completed_in_block)
        if warning and not messagebox.askyesno(_("Different durations in the block"), warning):
            log.info(
                "block_detail.mark_batch_complete.cancelled_on_duration_warning",
                day=self.day_num,
                group=self.group_name,
            )
            return

        master = self.master  # root Tk — sobrevive ao destroy deste Toplevel
        day_num = self.day_num
        group_name = self.group_name

        def _worker() -> None:
            try:
                all_data, unified_df, parsed_summary_files = self._build_partial_report_dataset(
                    summary_files
                )
                reports_dir = Path(self.project_manager.project_path) / "partial_reports"
                reports_dir.mkdir(parents=True, exist_ok=True)
                (
                    excel_output_name,
                    excel_output_path,
                    word_output_name,
                    _word_output_path,
                    write_fallback_used,
                ) = self._write_partial_report_outputs(
                    reports_dir,
                    parsed_summary_files,
                    all_data,
                    unified_df,
                )

                persisted = self.live_batch_coordinator.mark_block_complete(
                    group_name,
                    day_num,
                    unified_excel=excel_output_path,
                    session_count=len(all_data),
                )

                log.info(
                    "block_detail.mark_batch_complete.success",
                    day=day_num,
                    group=group_name,
                    excel_output=str(excel_output_path),
                    session_count=len(all_data),
                    persisted=persisted,
                )

                def _on_done() -> None:
                    # _publish_project_views_refresh só publica eventos via
                    # coordinators (não toca Tk), seguro após o destroy.
                    self._publish_project_views_refresh(
                        _("Batch complete: {block}").format(block=_block_label(day_num, group_name))
                    )
                    message = _(
                        "Batch '{block}' marked as complete.\n\n"
                        "📊 Excel: {excel}\n"
                        "📝 Word: {word}\n"
                        "🐟 {count} aggregated sessions\n\n"
                        "Reports in: {folder}"
                    ).format(
                        block=_block_label(day_num, group_name),
                        excel=excel_output_name,
                        word=word_output_name,
                        count=len(all_data),
                        folder=reports_dir,
                    )
                    if write_fallback_used:
                        message += _(
                            "\n\n⚠️ The default file was in use; the reports were "
                            "saved with a date/time suffix."
                        )
                    if not persisted:
                        message += _(
                            "\n\n⚠️ Could not record the completion in the project; check the log."
                        )
                    messagebox.showinfo(_("Batch Complete"), message, parent=master)

                master.after(0, _on_done)
            # except Exception justified: pipeline pandas/docx em thread de
            # fundo — qualquer falha deve virar feedback honesto na UI.
            except Exception as e:
                # O Python apaga ``e`` ao sair do except; a closure roda
                # depois (via after), então captura o texto agora.
                error_text = str(e)
                log.error(
                    "block_detail.mark_batch_complete.worker_failed",
                    day=day_num,
                    group=group_name,
                    error=error_text,
                    exc_info=True,
                )

                def _on_error() -> None:
                    messagebox.showerror(
                        _("Error — Batch not completed"),
                        _(
                            "Failed to generate the report for batch "
                            "'{block}':\n{error}\n\n"
                            "The batch was NOT marked as complete."
                        ).format(block=_block_label(day_num, group_name), error=error_text),
                        parent=master,
                    )

                master.after(0, _on_error)

        threading.Thread(target=_worker, name="MarkBatchComplete", daemon=True).start()
        messagebox.showinfo(
            _("Batch processing"),
            _(
                "Batch '{block}': the consolidated report (Excel + Word) is being "
                "generated in the background.\n\n"
                "You will be notified when it finishes."
            ).format(block=_block_label(self.day_num, self.group_name)),
        )
        self.destroy()
