"""
Extended unit tests for project_reporter in analysis/reporters/project_reporter.py.
"""

from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock, patch

import pandas as pd
from docx import Document

from zebtrack.analysis.reporters.project_reporter import (
    _render_comparative_boxplots,
    _render_descriptive_stats_table,
    export_multi_aquarium_reports,
    export_project_report,
)


def _doc_text(path: Path | str) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


class TestProjectReporterExtended:
    """Test project report generation and multi-aquarium export."""

    def test_export_project_report_with_detector_params_types(self, tmp_path: Path):
        df = pd.DataFrame(
            {
                "group_id": ["G1", "G2"],
                "video_duration_s": [60.0, 60.0],
            }
        )
        params = {
            "use_gpu": True,
            "confidence_threshold": 0.45,
            "model_name": "yolov8n",
        }
        out = tmp_path / "report_params"
        export_project_report(df, out, detector_params=params)

        docx_path = tmp_path / "report_params.docx"
        assert docx_path.exists()
        text = _doc_text(docx_path)
        assert "Detection Parameters" in text
        assert "Yes" in text  # boolean True
        assert "0.45" in text  # float format
        assert "yolov8n" in text

    def test_render_descriptive_stats_table(self):
        doc = Document()
        stats_df = pd.DataFrame(
            {
                "mean_speed_cm_s": [3.5, 4.2],
                "total_distance_cm": [100.0, 150.0],
            },
            index=["Control_Day1", "Treatment_Day1"],
        )
        _render_descriptive_stats_table(doc, stats_df)

        assert len(doc.tables) == 1
        table = doc.tables[0]
        # Header + 2 data rows
        assert len(table.rows) == 3
        assert table.rows[0].cells[0].text == "Group / Day"
        assert table.rows[1].cells[0].text == "Control_Day1"

    def test_render_comparative_boxplots(self):
        import matplotlib.pyplot as plt

        doc = Document()
        df = pd.DataFrame(
            {
                "group_id": ["A", "A", "B", "B"],
                "total_distance_cm": [10.0, 12.0, 20.0, 22.0],
            }
        )

        fig, ax = plt.subplots()
        ax.boxplot([[10.0, 12.0], [20.0, 22.0]])
        with patch(
            "zebtrack.analysis.visualization_generator.VisualizationGenerator.generate_comparative_boxplot",
            return_value=fig,
        ):
            _render_comparative_boxplots(doc, df)
        assert len(doc.paragraphs) > 0

    def test_export_multi_aquarium_reports_with_configs(self, tmp_path: Path):
        aq0_out = tmp_path / "aq0"
        aq1_out = tmp_path / "aq1"

        mock_res0 = MagicMock()
        mock_res1 = MagicMock()

        config0 = MagicMock(aquarium_id=0, group="Control", subject_id="1")
        config1 = MagicMock(aquarium_id=1, group="Treatment", subject_id="2")

        with patch("zebtrack.analysis.reporters.reporter_context.ReporterContext.from_analysis"):
            with patch("zebtrack.analysis.reporters.excel_reporter.ExcelReporter.export_summary"):
                with patch(
                    "zebtrack.analysis.reporters.word_reporter.WordReporter.export_individual_report"
                ):
                    res = export_multi_aquarium_reports(
                        results_by_aquarium={0: mock_res0, 1: mock_res1},
                        output_dirs_by_aquarium={0: aq0_out, 1: aq1_out},
                        base_name="session1",
                        aquarium_configs=[config0, config1],
                    )

                    assert 0 in res
                    assert 1 in res
                    assert "session1_Control_1_summary.xlsx" in res[0]["summary_path"]
                    assert "session1_Treatment_2_report.docx" in res[1]["report_path"]
