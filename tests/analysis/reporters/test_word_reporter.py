"""Structural + golden tests for the :class:`WordReporter` export pipeline.

The individual ``.docx`` is the per-subject artefact that ends up in the paper,
yet only its geotaxis-gating helper was previously exercised. These tests drive
the whole pipeline (``export_individual_report`` and the seven section builders),
re-open the produced document with ``python-docx`` and assert on its structure
and values -- the same reopen-and-read pattern used by
``tests/analysis/reporters/test_project_reporter.py``.

Assertions about headings/labels go through the reporter's own ``_`` translator
so they stay correct regardless of the machine locale (pt_BR vs C/en).

The matplotlib backend is left to the (headless) test environment, matching the
existing reporter tests which already embed rendered plots into ``.docx``.
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd
import pytest
from docx import Document

from zebtrack.analysis.reporters.reporter_context import _
from zebtrack.analysis.reporters.word_reporter import (
    WordReporter,
    _format_time_minutes_seconds,
)


def _doc_text(path: Path) -> str:
    """Return all paragraph + table text from a ``.docx`` file."""
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _find_table_by_header(path: Path, header_first_cell: str):
    """Return the first table whose top-left cell matches ``header_first_cell``."""
    doc = Document(str(path))
    for table in doc.tables:
        if table.rows and table.cell(0, 0).text == header_first_cell:
            return table
    return None


class TestExportIndividualReport:
    """End-to-end export covering the full builder chain."""

    def test_creates_docx_with_all_sections(self, reporter_ctx, tmp_path: Path) -> None:
        out = tmp_path / "report.docx"
        WordReporter(reporter_ctx).export_individual_report(out)

        assert out.exists()
        text = _doc_text(out)
        # Headings added unconditionally by the builder chain (locale-aware).
        for heading in (
            _("Experiment Metadata"),
            _("Metrics Summary"),
            _("ROI Reference Map"),
            _("Visualizations"),
            _("Appendix: Trajectory Validation"),
            _("Validation Details"),
        ):
            assert heading in text, f"missing section heading: {heading!r}"

    def test_appends_docx_suffix_when_missing(self, reporter_ctx, tmp_path: Path) -> None:
        # No ``.docx`` suffix -> the reporter must add it.
        out = tmp_path / "report_no_suffix"
        WordReporter(reporter_ctx).export_individual_report(out)
        assert (tmp_path / "report_no_suffix.docx").exists()

    def test_metadata_values_rendered(self, reporter_ctx, tmp_path: Path) -> None:
        out = tmp_path / "report.docx"
        WordReporter(reporter_ctx).export_individual_report(out)

        text = _doc_text(out)
        # The default fixture metadata is {"experiment_id": "test_001", ...}.
        assert "test_001" in text

    def test_metrics_table_is_populated(self, reporter_ctx, tmp_path: Path) -> None:
        out = tmp_path / "report.docx"
        WordReporter(reporter_ctx).export_individual_report(out)

        table = _find_table_by_header(out, _("Metric"))
        assert table is not None, "Metrics Summary table not found"
        # Header row + at least one metric row populated from the context.
        assert len(table.rows) > 1
        assert table.cell(0, 1).text == _("Value")

    def test_roi_event_log_present_with_rois(self, reporter_ctx, tmp_path: Path) -> None:
        out = tmp_path / "report.docx"
        WordReporter(reporter_ctx).export_individual_report(out)
        assert _("Appendix: ROI Event Log") in _doc_text(out)

    def test_roi_event_log_absent_without_rois(self, reporter_ctx_no_rois, tmp_path: Path) -> None:
        # ``r_analyzer is None`` -> the ROI event log appendix is skipped.
        out = tmp_path / "report.docx"
        WordReporter(reporter_ctx_no_rois).export_individual_report(out)
        text = _doc_text(out)
        assert _("Appendix: ROI Event Log") not in text
        # ...but the rest of the report (validation appendix) still renders.
        assert _("Appendix: Trajectory Validation") in text


class TestGeotaxisVisualizationBranch:
    """Both branches of the geotaxis ``if`` in ``_append_visualizations``."""

    def test_export_with_geotaxis_included(self, reporter_ctx, tmp_path: Path) -> None:
        reporter_ctx.behavioral_config = {"aquarium_perspective": "lateral"}
        assert WordReporter(reporter_ctx)._should_include_geotaxis_visualization() is True

        out = tmp_path / "report_lateral.docx"
        WordReporter(reporter_ctx).export_individual_report(out)
        assert out.exists()
        assert _("Visualizations") in _doc_text(out)

    def test_export_with_geotaxis_excluded(self, reporter_ctx, tmp_path: Path) -> None:
        reporter_ctx.behavioral_config = {"aquarium_perspective": "top_down"}
        assert WordReporter(reporter_ctx)._should_include_geotaxis_visualization() is False

        out = tmp_path / "report_topdown.docx"
        WordReporter(reporter_ctx).export_individual_report(out)
        assert out.exists()
        assert _("Visualizations") in _doc_text(out)


class TestProgressCallback:
    """The step-by-step variant must report bounded progress ending at 1.0.

    Note: the reporter's progress fractions are *not* strictly monotonic -- the
    visualisation section emits ``(4 + i) / total_steps`` (reaching ``10/11``
    when geotaxis adds a 7th plot) while the subsequent ROI-event-log step is
    hard-coded to ``9/11``, so the bar can step slightly backwards. That is a
    cosmetic UI quirk, not a data-correctness issue, so we only assert the
    contract callers actually depend on: bounded in [0, 1], multi-step, and a
    terminal 1.0.
    """

    def test_progress_is_bounded_and_completes(self, reporter_ctx, tmp_path: Path) -> None:
        calls: list[tuple[float, str]] = []

        out = tmp_path / "report.docx"
        WordReporter(reporter_ctx).export_individual_report_step_by_step(
            out, lambda p, s: calls.append((p, s))
        )

        assert calls, "progress_callback was never called"
        fractions = [p for p, _status in calls]
        assert all(0.0 <= p <= 1.0 for p in fractions)
        assert fractions[-1] == 1.0, "final progress is not 1.0"
        assert len(set(fractions)) > 1, "progress never advanced through steps"
        # Every status is a non-empty string.
        assert all(isinstance(s, str) and s for _p, s in calls)


class TestFormatTimeMinutesSeconds:
    """Golden cases for ``_format_time_minutes_seconds`` (minutes not padded)."""

    @pytest.mark.parametrize(
        ("seconds", "expected"),
        [
            (0, "0s"),
            (5, "5s"),
            (59, "59s"),
            (59.9, "59s"),
            (60, "1:00"),
            (65, "1:05"),
            (125, "2:05"),
            (3599, "59:59"),
            (3600, "1:00:00"),
            (3725, "1:02:05"),
            (7325, "2:02:05"),
        ],
    )
    def test_known_values(self, seconds: float, expected: str) -> None:
        assert _format_time_minutes_seconds(seconds) == expected

    def test_none_returns_na(self) -> None:
        assert _format_time_minutes_seconds(None) == "N/A"

    def test_nan_returns_na(self) -> None:
        assert _format_time_minutes_seconds(float("nan")) == "N/A"
        assert _format_time_minutes_seconds(pd.NA) == "N/A"  # type: ignore[arg-type]
