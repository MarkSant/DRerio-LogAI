"""Unit tests for ``zebtrack.analysis.reporters.project_reporter``.

Covers the project-level (aggregated) Word report and the multi-aquarium report
fan-out. ``export_project_report`` is exercised with a small aggregated frame
whose only metric is *not* in the boxplot list, so the matplotlib rendering path
is skipped (keeping the test fast and headless-safe). The generated ``.docx`` is
reopened with ``python-docx`` to assert the expected sections are present.

For ``export_multi_aquarium_reports`` the guard/skip branches are pinned (they
require no heavy ReporterContext).
"""

from __future__ import annotations

import pandas as pd
import pytest
from docx import Document

from zebtrack.analysis.reporters.project_reporter import (
    export_multi_aquarium_reports,
    export_project_report,
)


def _doc_text(path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


@pytest.fixture
def aggregated_df():
    # ``video_duration_s`` is a descriptive metric but NOT in the boxplot list,
    # so no matplotlib figures are generated.
    return pd.DataFrame(
        {
            "group_id": ["Control", "Control", "Treatment"],
            "video_duration_s": [120.0, 130.0, 110.0],
        }
    )


class TestExportProjectReport:
    def test_creates_docx_with_heading(self, aggregated_df, tmp_path):
        out = tmp_path / "project_report"
        export_project_report(aggregated_df, out)
        docx_path = tmp_path / "project_report.docx"
        assert docx_path.exists()
        assert "Aggregated Project Report" in _doc_text(docx_path)

    def test_descriptive_stats_section_present(self, aggregated_df, tmp_path):
        out = tmp_path / "r"
        export_project_report(aggregated_df, out)
        assert "Descriptive Statistics by Group" in _doc_text(tmp_path / "r.docx")

    def test_roi_color_legend_when_provided(self, aggregated_df, tmp_path):
        out = tmp_path / "r"
        export_project_report(aggregated_df, out, roi_colors={"Center": (255, 0, 0)})
        text = _doc_text(tmp_path / "r.docx")
        assert "ROI Color Legend" in text
        assert "Center" in text

    def test_detection_parameters_when_provided(self, aggregated_df, tmp_path):
        out = tmp_path / "r"
        export_project_report(aggregated_df, out, detector_params={"track_threshold": 0.25})
        assert "Detection Parameters" in _doc_text(tmp_path / "r.docx")

    def test_no_metrics_available_message(self, tmp_path):
        df = pd.DataFrame({"group_id": ["A", "B"]})  # no recognised metrics
        out = tmp_path / "r"
        export_project_report(df, out)
        assert "No metrics available" in _doc_text(tmp_path / "r.docx")


class TestExportMultiAquariumReports:
    def test_none_result_is_skipped(self, tmp_path):
        result = export_multi_aquarium_reports(
            results_by_aquarium={0: None},
            output_dirs_by_aquarium={0: tmp_path},
            base_name="clip",
        )
        assert result == {}

    def test_missing_output_dir_is_skipped(self, tmp_path):
        result = export_multi_aquarium_reports(
            results_by_aquarium={0: object()},  # truthy result...
            output_dirs_by_aquarium={},  # ...but no output dir → skip
            base_name="clip",
        )
        assert result == {}
