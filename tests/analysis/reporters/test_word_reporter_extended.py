"""
Extended unit tests for WordReporter.

Tests formatting helpers, geotaxis gating logic, per-animal ROI breakdowns,
unobserved time reporting, and validation statistics rendering.
"""

from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock

from docx import Document

from zebtrack.analysis.reporters.word_reporter import (
    WordReporter,
    _format_seconds_metric,
    _format_time_minutes_seconds,
)


def test_format_helpers():
    assert _format_seconds_metric(None) == "N/A"
    assert _format_seconds_metric(12.3456) == "12.35 s"
    assert _format_seconds_metric(0) == "0.00 s"

    assert _format_time_minutes_seconds(65.5) == "1:05"
    assert _format_time_minutes_seconds(10.0) == "10s"
    assert _format_time_minutes_seconds(None) == "N/A"


class TestWordReporterExtended:
    """Test extended WordReporter section generation."""

    def test_should_include_geotaxis_visualization(self):
        mock_ctx = MagicMock()
        mock_ctx._normalize_aquarium_perspective.return_value = "top_down"
        mock_ctx.behavioral_config = {"aquarium_perspective": "top_down"}
        reporter = WordReporter(mock_ctx)
        assert reporter._should_include_geotaxis_visualization() is False

        mock_ctx._normalize_aquarium_perspective.return_value = "lateral"
        mock_ctx.behavioral_config = {"aquarium_perspective": "lateral", "geotaxis_enabled": False}
        assert reporter._should_include_geotaxis_visualization() is False

        mock_ctx.behavioral_config = {"aquarium_perspective": "lateral", "geotaxis_enabled": True}
        assert reporter._should_include_geotaxis_visualization() is True

    def test_append_roi_coverage_and_per_animal_multiple_tracks(self, tmp_path: Path):
        doc = Document()
        mock_ctx = MagicMock()
        mock_ctx.report = {
            "analise_roi": {
                "tempo_nao_observado_s": 4.5,
                "por_animal": {
                    0: {
                        "tempo_gasto_por_roi": {"ROI_1": {"seconds": 10.5}},
                        "contagem_entradas": {"ROI_1": 3},
                        "distancia_por_roi": {"ROI_1": 55.2},
                        "tempo_nao_observado_s": 1.2,
                    },
                    1: {
                        "tempo_gasto_por_roi": {"ROI_1": {"seconds": 8.0}},
                        "contagem_entradas": {"ROI_1": 2},
                        "distancia_por_roi": {"ROI_1": 40.1},
                        "tempo_nao_observado_s": 3.3,
                    },
                },
            }
        }
        reporter = WordReporter(mock_ctx)
        reporter._append_roi_coverage_and_per_animal(doc)

        out_path = tmp_path / "coverage_test.docx"
        doc.save(str(out_path))

        saved_doc = Document(str(out_path))
        # Verify table created with per-animal rows
        assert len(saved_doc.tables) == 1
        table = saved_doc.tables[0]
        assert len(table.rows) == 3  # 1 header + 2 animals
        assert table.rows[1].cells[0].text == "0"
        assert table.rows[2].cells[0].text == "1"

    def test_append_validation_warnings_and_stats(self, tmp_path: Path):
        doc = Document()
        mock_ctx = MagicMock()
        mock_ctx.validation_stats = {
            "total_points": 1000,
            "interpolated_points": 50,
            "interpolation_percentage": 5.0,
            "max_consecutive_interpolated": 3,
            "out_of_bounds_points": 0,
            "excessive_speed_points": 2,
            "max_speed_cm_s": 45.2,
        }
        mock_ctx.validation_warnings = [
            "High speed detected around frame 120",
        ]

        reporter = WordReporter(mock_ctx)
        callback = MagicMock()
        reporter._append_validation_warnings(doc, callback, total_steps=10)

        out_path = tmp_path / "validation_test.docx"
        doc.save(str(out_path))

        saved_doc = Document(str(out_path))
        text = "\n".join(p.text for p in saved_doc.paragraphs)
        assert "Appendix: Trajectory Validation" in text
        assert "High speed detected around frame 120" in text
        callback.assert_called_once()
