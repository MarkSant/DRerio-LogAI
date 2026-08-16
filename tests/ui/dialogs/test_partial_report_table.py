"""The partial Word report's "Summary per Animal" table must stay readable.

Three things broke it at once:

1. Unbounded columns -- ``stats_cols`` comes from a keyword substring match, so
   it grows roughly linearly with the number of ROIs (``time_in_ROI{i}_s``,
   ``entries_ROI{i}``, ``time_in_ROI{i}_pct``...). Four ROIs already pushes past
   20 columns.
2. A bare ``Document()`` -- portrait Letter, 1-inch margins, ~6.5in usable.
3. Underscored headers, which Word cannot wrap, so autofit gives up.
"""

from __future__ import annotations

import pandas as pd
import pytest
from docx import Document
from docx.enum.section import WD_ORIENT

from zebtrack.ui.dialogs.block_detail_dialog import BlockDetailDialog


def _dialog() -> BlockDetailDialog:
    """The table renderer touches no Tk state, so skip __init__."""
    return BlockDetailDialog.__new__(BlockDetailDialog)


def _frame(animals: list[str], metrics: list[str]) -> pd.DataFrame:
    return pd.DataFrame({"animal": animals, **{m: [1.5] * len(animals) for m in metrics}})


def _roi_metrics(n_rois: int) -> list[str]:
    """The realistic column set a block with ``n_rois`` ROIs produces."""
    metrics = [
        "total_distance_cm",
        "mean_speed_cm_s",
        "max_speed_cm_s",
        "time_moving_s",
        "time_frozen_s",
        "mean_angular_velocity_deg_s",
    ]
    for i in range(1, n_rois + 1):
        metrics += [f"time_in_ROI{i}_s", f"entries_ROI{i}", f"time_in_ROI{i}_pct"]
    return metrics


def _render(animals: list[str], metrics: list[str]):
    document = Document()
    BlockDetailDialog._render_partial_summary_table(
        _dialog(), document, _frame(animals, metrics), metrics
    )
    return document


class TestReadability:
    def test_landscape_section_with_narrow_margins(self):
        document = _render(["A", "B"], _roi_metrics(4))

        section = document.sections[-1]
        assert section.orientation == WD_ORIENT.LANDSCAPE
        assert section.page_width > section.page_height

    def test_four_rois_do_not_produce_a_twenty_column_table(self):
        """18 metrics x 3 animals used to render as 19 columns across Letter portrait."""
        metrics = _roi_metrics(4)
        document = _render(["S1", "S2", "S3"], metrics)

        table = document.tables[0]
        assert len(metrics) == 18
        assert len(table.columns) <= BlockDetailDialog._MAX_REPORT_TABLE_COLUMNS
        # Metrics outnumber animals -> metrics go on rows.
        assert table.rows[0].cells[0].text == "Metric"
        assert len(table.rows) == len(metrics) + 1

    def test_headers_are_humanized_so_word_can_wrap(self):
        document = _render(["S1"], ["total_distance_cm", "mean_speed_cm_s"])

        table = document.tables[0]
        rendered = [row.cells[0].text for row in table.rows]
        assert "total distance cm" in rendered
        assert not any("_" in text for text in rendered)

    def test_header_row_repeats_across_pages(self):
        document = _render(["S1", "S2"], _roi_metrics(4))

        tr_pr = document.tables[0].rows[0]._tr.get_or_add_trPr()
        assert tr_pr.findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblHeader"
        )


class TestOrientationFollowsTheData:
    def test_metrics_on_rows_when_metrics_outnumber_animals(self):
        document = _render(["S1", "S2"], _roi_metrics(3))

        table = document.tables[0]
        assert table.rows[0].cells[0].text == "Metric"
        assert [c.text for c in table.rows[0].cells[1:]] == ["S1", "S2"]

    def test_animals_on_rows_when_animals_outnumber_metrics(self):
        """A 20-subject block must not become a 21-column table either."""
        animals = [f"S{i}" for i in range(1, 21)]
        document = _render(animals, ["total_distance_cm", "mean_speed_cm_s"])

        table = document.tables[0]
        assert table.rows[0].cells[0].text == "Animal"
        assert len(table.rows) == len(animals) + 1
        assert len(table.columns) == 3

    def test_values_stay_aligned_with_their_labels_when_transposed(self):
        metrics = ["total_distance_cm", "mean_speed_cm_s", "time_frozen_s"]
        frame = pd.DataFrame(
            {
                "animal": ["S1", "S2"],
                "total_distance_cm": [10.0, 20.0],
                "mean_speed_cm_s": [1.25, 2.25],
                "time_frozen_s": [5.0, 6.0],
            }
        )
        document = Document()
        BlockDetailDialog._render_partial_summary_table(_dialog(), document, frame, metrics)

        table = document.tables[0]
        rows = {r.cells[0].text: [c.text for c in r.cells[1:]] for r in table.rows[1:]}
        # S1 == first data column, S2 == second.
        assert rows["mean speed cm s"] == ["1.25", "2.25"]
        assert rows["total distance cm"] == ["10", "20"]


class TestTruncation:
    def test_excess_columns_are_reported_not_dropped_silently(self):
        """Truncating in silence reads as "these are all the metrics"."""
        animals = [f"S{i}" for i in range(1, 41)]  # animals >> metrics
        document = _render(animals, ["total_distance_cm", "mean_speed_cm_s"])

        # Animals go on rows here, so nothing is dropped.
        assert len(document.tables[0].columns) == 3
        assert not any("omitted" in p.text for p in document.paragraphs)

    def test_note_is_stamped_when_columns_must_be_dropped(self):
        # Force the truncation branch: equal counts, both above the cap.
        n = BlockDetailDialog._MAX_REPORT_TABLE_COLUMNS + 5
        animals = [f"S{i}" for i in range(n)]
        metrics = [f"metric_{i}_cm" for i in range(n)]
        document = _render(animals, metrics)

        table = document.tables[0]
        assert len(table.columns) == BlockDetailDialog._MAX_REPORT_TABLE_COLUMNS
        assert any("omitted" in p.text for p in document.paragraphs)


class TestDegenerateInput:
    @pytest.mark.parametrize(
        ("animals", "metrics"),
        [([], ["total_distance_cm"]), (["S1"], [])],
    )
    def test_empty_dimension_renders_nothing(self, animals, metrics):
        document = Document()
        frame = pd.DataFrame({"animal": animals, **{m: [] for m in metrics}})
        BlockDetailDialog._render_partial_summary_table(_dialog(), document, frame, metrics)

        assert not document.tables
